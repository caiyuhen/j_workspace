"""
Word (.docx) 文档导出模块 (美化版)
Document Exporter Module

支持:
- 医学论文导出 (IMRaD 结构)
- 基金申请书导出
- Response Letter 导出
- RCT 试验方案导出
"""

import os
from typing import Dict, List, Any, Optional
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

from .styles import Colors, Fonts, Spacing, BRAND_NAME, BRAND_TAGLINE


class BaseDocumentExporter:
    """Word 文档导出基类 - 美化版"""

    def __init__(self):
        if not HAS_DOCX:
            raise ImportError("python-docx is required. Install: pip install python-docx")
        self.doc = Document()
        self._setup_styles()
        self._setup_page()

    def _setup_page(self):
        """设置页面边距"""
        sections = self.doc.sections[0]
        sections.top_margin = Cm(2.54)
        sections.bottom_margin = Cm(2.54)
        sections.left_margin = Cm(3.17)
        sections.right_margin = Cm(3.17)

    def _setup_styles(self):
        """设置统一文档样式"""
        # 正文样式
        style = self.doc.styles['Normal']
        font = style.font
        font.name = Fonts.LATIN_BODY
        font.size = Fonts.SIZE_BODY
        font.color.rgb = Colors.TEXT
        style.element.rPr.rFonts.set(qn('w:eastAsia'), Fonts.CJK_BODY)
        pf = style.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        pf.space_before = Spacing.PARA_BEFORE
        pf.space_after = Spacing.PARA_AFTER

        # 标题样式 (Heading 1-3)
        heading_configs = [
            (1, Fonts.SIZE_H1, Colors.PRIMARY_DK),
            (2, Fonts.SIZE_H2, Colors.PRIMARY),
            (3, Fonts.SIZE_H3, Colors.PRIMARY),
        ]
        for level, size, color in heading_configs:
            heading = self.doc.styles[f'Heading {level}']
            heading.font.name = Fonts.LATIN_HEAD
            heading.font.size = size
            heading.font.bold = True
            heading.font.color.rgb = color
            heading.element.rPr.rFonts.set(qn('w:eastAsia'), Fonts.CJK_HEADING)
            hpf = heading.paragraph_format
            hpf.space_before = Spacing.HEADING_BEFORE
            hpf.space_after = Spacing.HEADING_AFTER
            hpf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    def add_cover_page(self, title: str, subtitle: str = "", 
                       author: str = "", doc_type: str = ""):
        """添加专业封面页"""
        # 顶部装饰条
        top_bar = self.doc.add_paragraph()
        top_bar.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = top_bar.add_run("━" * 40)
        run.font.size = Pt(8)
        run.font.color.rgb = Colors.PRIMARY
        top_bar.paragraph_format.space_after = Pt(60)

        # 品牌标识
        brand_p = self.doc.add_paragraph()
        brand_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = brand_p.add_run(BRAND_NAME.upper())
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = Colors.PRIMARY
        run.font.name = Fonts.LATIN_HEAD
        run.element.rPr.rFonts.set(qn('w:eastAsia'), Fonts.CJK_HEADING)

        tag_p = self.doc.add_paragraph()
        tag_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = tag_p.add_run(BRAND_TAGLINE)
        run.font.size = Pt(10)
        run.font.color.rgb = Colors.TEXT_LT
        run.font.italic = True
        tag_p.paragraph_format.space_after = Pt(40)

        # 文档类型标签
        if doc_type:
            type_p = self.doc.add_paragraph()
            type_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = type_p.add_run(doc_type)
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = Colors.WHITE
            run.font.name = Fonts.LATIN_HEAD
            # 添加浅蓝背景效果通过段落底纹（简化处理）
            type_p.paragraph_format.space_after = Pt(30)

        # 主标题
        title_p = self.doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_p.add_run(title)
        run.font.size = Fonts.SIZE_TITLE
        run.font.bold = True
        run.font.color.rgb = Colors.PRIMARY_DK
        run.font.name = Fonts.LATIN_HEAD
        run.element.rPr.rFonts.set(qn('w:eastAsia'), Fonts.CJK_HEADING)
        title_p.paragraph_format.space_after = Pt(12)

        # 副标题
        if subtitle:
            sub_p = self.doc.add_paragraph()
            sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = sub_p.add_run(subtitle)
            run.font.size = Fonts.SIZE_H3
            run.font.color.rgb = Colors.TEXT_LT
            sub_p.paragraph_format.space_after = Pt(60)

        # 作者
        if author:
            auth_p = self.doc.add_paragraph()
            auth_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = auth_p.add_run(author)
            run.font.size = Fonts.SIZE_BODY
            run.font.color.rgb = Colors.TEXT
            auth_p.paragraph_format.space_after = Pt(12)

        # 日期
        date_p = self.doc.add_paragraph()
        date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = date_p.add_run(datetime.now().strftime("%Y-%m-%d"))
        run.font.size = Fonts.SIZE_SMALL
        run.font.color.rgb = Colors.TEXT_LT
        date_p.paragraph_format.space_after = Pt(80)

        # 底部装饰条
        bot_bar = self.doc.add_paragraph()
        bot_bar.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = bot_bar.add_run("━" * 40)
        run.font.size = Pt(8)
        run.font.color.rgb = Colors.PRIMARY

        # 分页
        self.doc.add_page_break()

    def add_header_footer(self, doc_type: str = ""):
        """添加页眉页脚"""
        section = self.doc.sections[0]

        # 页眉
        header = section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.text = f"{BRAND_NAME}  |  {doc_type}" if doc_type else BRAND_NAME
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in header_para.runs:
            run.font.size = Fonts.SIZE_CAPTION
            run.font.color.rgb = Colors.TEXT_LT
            run.font.name = Fonts.LATIN_BODY

        # 页脚 - 页码
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_para.add_run("— ")
        run.font.size = Fonts.SIZE_CAPTION
        run.font.color.rgb = Colors.TEXT_LT
        # 添加页码字段
        run2 = footer_para.add_run()
        run2.font.size = Fonts.SIZE_CAPTION
        run2.font.color.rgb = Colors.TEXT_LT
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run2._r.append(fldChar1)
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run2._r.append(instrText)
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run2._r.append(fldChar2)
        run3 = footer_para.add_run(" —")
        run3.font.size = Fonts.SIZE_CAPTION
        run3.font.color.rgb = Colors.TEXT_LT

    def add_title(self, title: str, level: int = 1):
        """添加标题"""
        self.doc.add_heading(title, level=level)

    def add_paragraph(self, text: str, bold: bool = False, italic: bool = False,
                      alignment=WD_ALIGN_PARAGRAPH.LEFT):
        """添加段落（美化版）"""
        p = self.doc.add_paragraph()
        p.alignment = alignment
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.name = Fonts.LATIN_BODY
        run.font.size = Fonts.SIZE_BODY
        run.font.color.rgb = Colors.TEXT
        run.element.rPr.rFonts.set(qn('w:eastAsia'), Fonts.CJK_BODY)
        return p

    def add_styled_paragraph(self, text: str, label: str = "", 
                             label_color=None, bold: bool = False):
        """添加带标签的段落（如 关键词: xxx）"""
        p = self.doc.add_paragraph()
        if label:
            run_label = p.add_run(f"{label}: ")
            run_label.bold = True
            run_label.font.color.rgb = label_color or Colors.PRIMARY
            run_label.font.size = Fonts.SIZE_BODY
            run_label.font.name = Fonts.LATIN_BODY
            run_label.element.rPr.rFonts.set(qn('w:eastAsia'), Fonts.CJK_BODY)
        run_text = p.add_run(text)
        run_text.bold = bold
        run_text.font.size = Fonts.SIZE_BODY
        run_text.font.name = Fonts.LATIN_BODY
        run_text.element.rPr.rFonts.set(qn('w:eastAsia'), Fonts.CJK_BODY)
        return p

    def add_table(self, headers: List[str], rows: List[List[str]],
                  col_widths: List[int] = None):
        """添加专业医学表格（美化版）"""
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        # 设置表格整体样式
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')

        # 表头行
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            cell = hdr_cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(header)
            run.bold = True
            run.font.size = Fonts.SIZE_SMALL
            run.font.color.rgb = Colors.WHITE
            run.font.name = Fonts.LATIN_BODY
            run.element.rPr.rFonts.set(qn('w:eastAsia'), Fonts.CJK_BODY)
            # 表头背景色
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1E5AA8" w:val="clear"/>')
            cell._tc.get_or_add_tcPr().append(shading)

        # 数据行（斑马纹）
        for row_idx, row_data in enumerate(rows):
            row_cells = table.add_row().cells
            bg_color = "F2F7FC" if row_idx % 2 == 1 else "FFFFFF"
            for i, cell_text in enumerate(row_data):
                cell = row_cells[i]
                cell.text = ""
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(str(cell_text))
                run.font.size = Fonts.SIZE_SMALL
                run.font.color.rgb = Colors.TEXT
                run.font.name = Fonts.LATIN_BODY
                run.element.rPr.rFonts.set(qn('w:eastAsia'), Fonts.CJK_BODY)
                # 斑马纹背景
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_color}" w:val="clear"/>')
                cell._tc.get_or_add_tcPr().append(shading)
                # 细边框
                tcPr = cell._tc.get_or_add_tcPr()
                tcBorders = parse_xml(
                    f'<w:tcBorders {nsdecls("w")}>'
                    f'<w:top w:val="single" w:sz="4" w:space="0" w:color="D0D7DE"/>'
                    f'<w:left w:val="single" w:sz="4" w:space="0" w:color="D0D7DE"/>'
                    f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="D0D7DE"/>'
                    f'<w:right w:val="single" w:sz="4" w:space="0" w:color="D0D7DE"/>'
                    f'</w:tcBorders>'
                )
                tcPr.append(tcBorders)

        # 设置列宽
        if col_widths:
            for row in table.rows:
                for idx, width in enumerate(col_widths):
                    if idx < len(row.cells):
                        row.cells[idx].width = Cm(width)

        return table

    def add_references(self, references: List[str]):
        """添加专业参考文献列表（悬挂缩进）"""
        if not references:
            return
        self.add_title("参考文献", level=2)
        for i, ref in enumerate(references, 1):
            p = self.doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.0)
            p.paragraph_format.first_line_indent = Cm(-1.0)
            p.paragraph_format.space_after = Pt(4)
            run_num = p.add_run(f"[{i}] ")
            run_num.font.size = Fonts.SIZE_SMALL
            run_num.font.color.rgb = Colors.PRIMARY
            run_num.font.name = Fonts.LATIN_BODY
            run_text = p.add_run(ref)
            run_text.font.size = Fonts.SIZE_SMALL
            run_text.font.color.rgb = Colors.TEXT
            run_text.font.name = Fonts.LATIN_BODY
            run_text.element.rPr.rFonts.set(qn('w:eastAsia'), Fonts.CJK_BODY)

    def save(self, file_path: str):
        """保存文档"""
        os.makedirs(os.path.dirname(file_path) or '.', exist_ok=True)
        self.doc.save(file_path)
        return file_path


class PaperExporter(BaseDocumentExporter):
    """医学论文导出器 (IMRaD 结构) - 美化版"""

    def export_paper(self, paper_data: Dict[str, Any], file_path: str) -> str:
        """
        将论文数据导出为 Word 文档

        Args:
            paper_data: 论文数据字典，包含 title, authors, abstract, keywords,
                       introduction, methods, results, discussion, conclusion, references
            file_path: 输出文件路径

        Returns:
            保存的文件路径
        """
        title = paper_data.get('title', 'Untitled Paper')
        authors = paper_data.get('authors', '')

        # 封面页
        self.add_cover_page(
            title=title,
            subtitle=paper_data.get('subtitle', ''),
            author=authors,
            doc_type="医学论文"
        )

        # 页眉页脚
        self.add_header_footer(doc_type="医学论文")

        # 作者信息
        if authors:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(authors)
            run.italic = True
            run.font.size = Fonts.SIZE_BODY
            run.font.color.rgb = Colors.TEXT_LT

        # 摘要
        abstract = paper_data.get('abstract', '')
        if abstract:
            self.add_title('摘要', level=2)
            self.add_paragraph(abstract)

        # 关键词
        keywords = paper_data.get('keywords', [])
        if keywords:
            self.add_styled_paragraph(
                ', '.join(keywords),
                label="关键词",
                label_color=Colors.PRIMARY
            )

        self.doc.add_page_break()

        # IMRaD 结构
        sections = [
            ('引言', 'introduction'),
            ('方法', 'methods'),
            ('结果', 'results'),
            ('讨论', 'discussion'),
            ('结论', 'conclusion'),
        ]

        for section_title, key in sections:
            content = paper_data.get(key, '')
            if content:
                self.add_title(section_title, level=2)
                self.add_paragraph(content)

        # 参考文献（使用美化版）
        references = paper_data.get('references', [])
        if references:
            self.add_references(references)

        return self.save(file_path)


class GrantProposalExporter(BaseDocumentExporter):
    """基金申请书导出器 - 美化版"""

    def export_proposal(self, proposal_data: Dict[str, Any],
                        file_path: str) -> str:
        """
        将基金申请书导出为 Word 文档

        Args:
            proposal_data: 申请书数据
            file_path: 输出文件路径
        """
        title = proposal_data.get('title', '基金申请书')
        applicant = proposal_data.get('applicant', '')
        institution = proposal_data.get('institution', '')
        author_info = f"{applicant}  |  {institution}" if applicant and institution else applicant or institution

        # 封面页
        self.add_cover_page(
            title=title,
            subtitle=proposal_data.get('subtitle', ''),
            author=author_info,
            doc_type="基金申请书"
        )

        # 页眉页脚
        self.add_header_footer(doc_type="基金申请书")

        # 元数据
        metadata = [
            ('申请类型', proposal_data.get('grant_type', '')),
            ('研究领域', proposal_data.get('research_area', '')),
            ('申请人', applicant),
            ('依托单位', institution),
            ('申请日期', proposal_data.get('date', datetime.now().strftime('%Y-%m-%d'))),
        ]

        for label, value in metadata:
            if value:
                self.add_styled_paragraph(str(value), label=label, label_color=Colors.PRIMARY)

        self.doc.add_page_break()

        # 正文各部分
        sections = [
            ('一、立项依据', 'rationale'),
            ('二、研究内容', 'research_content'),
            ('三、研究目标', 'objectives'),
            ('四、拟解决的关键科学问题', 'key_problems'),
            ('五、研究方案与技术路线', 'methodology'),
            ('六、可行性分析', 'feasibility'),
            ('七、特色与创新之处', 'innovation'),
            ('八、年度计划', 'timeline'),
            ('九、预期成果', 'expected_outcomes'),
        ]

        for section_title, key in sections:
            content = proposal_data.get(key, '')
            if content:
                self.add_title(section_title, level=2)
                self.add_paragraph(content)

        # 经费预算表
        budget = proposal_data.get('budget', {})
        if budget:
            self.add_title('十、经费预算', level=2)
            headers = ['科目', '金额（万元）', '占比', '说明']
            rows = []
            total = budget.get('total', 0)
            for item in budget.get('items', []):
                name = item.get('name', '')
                amount = item.get('amount', 0)
                pct = f"{amount / total * 100:.1f}%" if total > 0 else '0%'
                rows.append([name, f"{amount:.2f}", pct, item.get('notes', '')])
            rows.append(['合计', f"{total:.2f}", '100%', ''])
            self.add_table(headers, rows)

        return self.save(file_path)


class ResponseLetterExporter(BaseDocumentExporter):
    """Response Letter 导出器 - 美化版"""

    def export_response_letter(self, letter_data: Dict[str, Any],
                               file_path: str) -> str:
        """
        将 Response Letter 导出为 Word 文档

        Args:
            letter_data: 包含 manuscript_id, title, authors, responses
            file_path: 输出文件路径
        """
        title = letter_data.get('title', 'Response to Reviewers')

        # 封面页
        self.add_cover_page(
            title="Response to Reviewers",
            subtitle=title,
            author=letter_data.get('authors', ''),
            doc_type="审稿回复信"
        )

        # 页眉页脚
        self.add_header_footer(doc_type="审稿回复信")

        # 稿件信息
        self.add_styled_paragraph(
            letter_data.get('manuscript_id', 'N/A'),
            label="Manuscript ID",
            label_color=Colors.PRIMARY
        )

        # 问候语
        greeting = (
            "Dear Editor and Reviewers,\n\n"
            "We would like to express our sincere gratitude for the constructive "
            "comments and suggestions. We have carefully revised the manuscript "
            "according to the reviewers' comments. The detailed point-by-point "
            "responses are provided below."
        )
        self.add_paragraph(greeting)

        # 逐条回复
        responses = letter_data.get('responses', [])
        for i, resp in enumerate(responses, 1):
            self.add_title(f"Comment {i}", level=2)

            self.add_styled_paragraph(
                resp.get('comment', ''),
                label="Reviewer's Comment",
                label_color=Colors.TEXT
            )

            self.add_styled_paragraph(
                resp.get('response', ''),
                label="Author's Response",
                label_color=Colors.PRIMARY
            )

            if resp.get('changes'):
                self.add_styled_paragraph(
                    resp['changes'],
                    label="Revisions Made",
                    label_color=Colors.ACCENT
                )

        return self.save(file_path)


class ProtocolExporter(BaseDocumentExporter):
    """RCT 试验方案导出器 - 美化版"""

    def export_protocol(self, protocol_data: Dict[str, Any],
                        file_path: str) -> str:
        """
        将 RCT 试验方案导出为 Word 文档

        Args:
            protocol_data: 试验方案数据
            file_path: 输出文件路径
        """
        study_info = protocol_data.get('study_info', {})
        title = study_info.get('title', '临床试验方案')
        indication = study_info.get('indication', '')

        # 封面页
        self.add_cover_page(
            title=title,
            subtitle=indication,
            author=study_info.get('sponsor', ''),
            doc_type="临床试验方案"
        )

        # 页眉页脚
        self.add_header_footer(doc_type="临床试验方案")

        # 基本信息
        self.add_title('一、研究基本信息', level=2)
        info_items = [
            ('研究类型', study_info.get('study_type', '')),
            ('试验分期', study_info.get('phase', '')),
            ('适应症', indication),
            ('研究时长', f"{study_info.get('duration_months', '')} 个月"),
        ]
        for label, value in info_items:
            if value:
                self.add_styled_paragraph(str(value), label=label, label_color=Colors.PRIMARY)

        # 研究目的
        objectives = protocol_data.get('study_objectives', {})
        if objectives:
            self.add_title('二、研究目的', level=2)
            for k, v in objectives.items():
                self.add_styled_paragraph(str(v), label=k, label_color=Colors.TEXT)

        # 终点
        endpoints = protocol_data.get('endpoints', {})
        if endpoints:
            self.add_title('三、研究终点', level=2)
            for k, v in endpoints.items():
                self.add_styled_paragraph(str(v), label=k, label_color=Colors.TEXT)

        # 入排标准
        inclusion = protocol_data.get('inclusion_criteria', [])
        exclusion = protocol_data.get('exclusion_criteria', [])
        if inclusion:
            self.add_title('四、入选标准', level=2)
            for item in inclusion:
                self.doc.add_paragraph(item, style='List Number')
        if exclusion:
            self.add_title('五、排除标准', level=2)
            for item in exclusion:
                self.doc.add_paragraph(item, style='List Number')

        # 样本量
        sample_size = protocol_data.get('sample_size', {})
        if sample_size:
            self.add_title('六、样本量', level=2)
            self.add_paragraph(str(sample_size))

        # 统计分析
        stats = protocol_data.get('statistical_analysis', {})
        if stats:
            self.add_title('七、统计分析计划', level=2)
            self.add_paragraph(str(stats))

        # 伦理
        ethics = protocol_data.get('ethical_considerations', {})
        if ethics:
            self.add_title('八、伦理考虑', level=2)
            for k, v in ethics.items():
                self.add_styled_paragraph(str(v), label=k, label_color=Colors.TEXT)

        return self.save(file_path)
