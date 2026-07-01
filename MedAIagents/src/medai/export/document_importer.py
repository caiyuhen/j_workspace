"""
Office 文档导入模块
Document Importer Module

支持:
- Word (.docx) 文本提取
- Excel (.xlsx) 数据读取
"""

import os
from typing import Dict, List, Any, Optional

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from openpyxl import load_workbook
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False


class WordImporter:
    """Word 文档导入器"""

    def __init__(self):
        if not HAS_DOCX:
            raise ImportError("python-docx is required. Install: pip install python-docx")

    def extract_text(self, file_path: str) -> str:
        """
        提取 Word 文档的全部文本

        Args:
            file_path: Word 文件路径

        Returns:
            文档纯文本
        """
        doc = DocxDocument(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n".join(paragraphs)

    def extract_structure(self, file_path: str) -> Dict[str, Any]:
        """
        提取 Word 文档的结构化内容

        Args:
            file_path: Word 文件路径

        Returns:
            结构化数据，包含 paragraphs, tables, headings
        """
        doc = DocxDocument(file_path)

        result = {
            "file_path": file_path,
            "paragraphs": [],
            "headings": [],
            "tables": [],
            "metadata": {
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
            }
        }

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style = para.style.name if para.style else "Normal"
            result["paragraphs"].append({
                "text": text,
                "style": style,
                "is_heading": style.startswith("Heading"),
                "heading_level": int(style.replace("Heading ", ""))
                if style.startswith("Heading") and style.replace("Heading ", "").isdigit()
                else 0,
            })
            if style.startswith("Heading"):
                result["headings"].append(text)

        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            result["tables"].append(table_data)

        return result

    def extract_imrad(self, file_path: str) -> Dict[str, str]:
        """
        从 Word 文档中提取 IMRaD 结构

        Args:
            file_path: Word 文件路径

        Returns:
            IMRaD 各部分文本
        """
        doc = DocxDocument(file_path)
        sections = {
            "title": "",
            "abstract": "",
            "introduction": "",
            "methods": "",
            "results": "",
            "discussion": "",
            "conclusion": "",
            "references": "",
        }

        current_section = None
        section_map = {
            "introduction": ["introduction", "intro", "背景", "前言", "引言"],
            "methods": ["methods", "method", "方法", "材料与方法", "methodology"],
            "results": ["results", "result", "结果"],
            "discussion": ["discussion", "discuss", "讨论"],
            "conclusion": ["conclusion", "conclusions", "结论"],
            "abstract": ["abstract", "摘要"],
            "references": ["references", "reference", "参考文献"],
        }

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style = para.style.name if para.style else ""
            text_lower = text.lower()

            # 检测标题
            if style == "Title" or (not sections["title"] and len(text) < 200):
                sections["title"] = text
                continue

            # 检测章节标题
            is_heading = style.startswith("Heading") or text.isupper()
            if is_heading:
                found = False
                for section_key, keywords in section_map.items():
                    for kw in keywords:
                        if kw in text_lower:
                            current_section = section_key
                            found = True
                            break
                    if found:
                        break
                continue

            # 累积正文
            if current_section and current_section in sections:
                sections[current_section] += text + "\n"

        return {k: v.strip() for k, v in sections.items()}


class ExcelImporter:
    """Excel 文档导入器"""

    def __init__(self):
        if not HAS_OPENPYXL:
            raise ImportError("openpyxl is required. Install: pip install openpyxl")

    def read_sheet(self, file_path: str, sheet_name: Optional[str] = None,
                   header_row: int = 1) -> List[Dict[str, Any]]:
        """
        读取 Excel 工作表为字典列表

        Args:
            file_path: Excel 文件路径
            sheet_name: 工作表名称，默认第一个
            header_row: 表头行号（1-based）

        Returns:
            字典列表
        """
        wb = load_workbook(file_path, data_only=True)
        ws = wb[sheet_name] if sheet_name else wb.active

        # 读取表头
        headers = []
        for cell in ws[header_row]:
            headers.append(str(cell.value) if cell.value else f"Column_{cell.column}")

        # 读取数据
        data = []
        for row in ws.iter_rows(min_row=header_row + 1, values_only=True):
            if all(v is None for v in row):
                continue
            row_dict = {}
            for i, header in enumerate(headers):
                row_dict[header] = row[i] if i < len(row) else None
            data.append(row_dict)

        return data

    def read_all_sheets(self, file_path: str) -> Dict[str, List[Dict[str, Any]]]:
        """
        读取 Excel 所有工作表

        Args:
            file_path: Excel 文件路径

        Returns:
            {sheet_name: [dict, ...]}
        """
        wb = load_workbook(file_path, data_only=True)
        result = {}
        for sheet_name in wb.sheetnames:
            result[sheet_name] = self.read_sheet(file_path, sheet_name)
        return result

    def read_survival_data(self, file_path: str,
                           sheet_name: Optional[str] = None) -> List[Dict[str, Any]]:
        """
        读取生存分析数据

        Args:
            file_path: Excel 文件路径
            sheet_name: 工作表名称

        Returns:
            生存记录列表
        """
        data = self.read_sheet(file_path, sheet_name)
        survival_records = []
        for row in data:
            record = {
                "patient_id": str(row.get("Patient ID", row.get("patient_id", ""))),
                "time": float(row.get("Time (months)", row.get("time", 0)) or 0),
                "event": int(row.get("Event (0=censored, 1=event)", row.get("event", 0)) or 0),
                "group": str(row.get("Group", row.get("group", ""))),
            }
            # 提取协变量
            covariates = {}
            for key in ["age", "stage", "gender", "treatment"]:
                if key in row and row[key] is not None:
                    try:
                        covariates[key] = float(row[key])
                    except (ValueError, TypeError):
                        covariates[key] = str(row[key])
            if covariates:
                record["covariates"] = covariates
            survival_records.append(record)
        return survival_records

    def read_journal_database(self, file_path: str,
                              sheet_name: Optional[str] = None) -> List[Dict[str, Any]]:
        """
        读取期刊数据库

        Args:
            file_path: Excel 文件路径
            sheet_name: 工作表名称

        Returns:
            期刊信息列表
        """
        data = self.read_sheet(file_path, sheet_name)
        journals = []
        for row in data:
            journal = {
                "name": str(row.get("Journal Name", row.get("name", ""))),
                "impact_factor": row.get("Impact Factor", row.get("impact_factor", "")),
                "jcr_quartile": str(row.get("JCR Quartile", row.get("jcr_quartile", ""))),
                "cas_quartile": str(row.get("CAS Quartile", row.get("cas_quartile", ""))),
                "field": str(row.get("Field", row.get("field", ""))),
                "oa_policy": str(row.get("OA Policy", row.get("oa_policy", ""))),
                "review_period": str(row.get("Review Period", row.get("review_period", ""))),
            }
            journals.append(journal)
        return journals
