"""
文档解析器
支持 PDF、DOCX、TXT、Markdown 等格式的文本提取和分块
"""

import os
import re
from typing import List, Dict, Any, Optional
from pathlib import Path
from loguru import logger


class DocumentParser:
    """文档解析器 — 提取文本并分块"""
    
    # 支持的文件扩展名
    SUPPORTED_EXTENSIONS = {'.txt', '.md', '.markdown', '.pdf', '.docx', '.doc'}
    
    # 分块大小（字符数）
    CHUNK_SIZE = 800
    CHUNK_OVERLAP = 100
    
    def __init__(self, chunk_size: int = 800, chunk_overlap: int = 100):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
    
    def parse_file(self, file_path: str) -> Dict[str, Any]:
        """解析文件，返回 {title, content, chunks, metadata}
        
        Args:
            file_path: 文件绝对路径
        
        Returns:
            解析结果字典
        """
        path = Path(file_path)
        if not path.exists():
            return {"success": False, "error": f"文件不存在: {file_path}"}
        
        ext = path.suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            return {"success": False, "error": f"不支持的文件格式: {ext}，支持 {self.SUPPORTED_EXTENSIONS}"}
        
        try:
            # 提取原始文本
            if ext == '.pdf':
                text = self._parse_pdf(file_path)
            elif ext in ('.docx', '.doc'):
                text = self._parse_docx(file_path)
            elif ext in ('.txt', '.md', '.markdown'):
                text = self._parse_text(file_path)
            else:
                text = self._parse_text(file_path)
            
            if text is None:
                return {"success": False, "error": "文本提取失败，可能缺少依赖库"}
            
            # 清理文本
            text = self._clean_text(text)
            
            # 分块
            chunks = self._chunk_text(text)
            
            return {
                "success": True,
                "title": path.stem,
                "content": text,
                "chunks": chunks,
                "metadata": {
                    "filename": path.name,
                    "extension": ext,
                    "size_bytes": path.stat().st_size,
                    "chunk_count": len(chunks),
                }
            }
        except Exception as e:
            logger.error(f"解析文件失败 {file_path}: {e}")
            return {"success": False, "error": str(e)}
    
    def _parse_pdf(self, file_path: str) -> Optional[str]:
        """解析 PDF 文件"""
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(file_path)
            texts = []
            for page in doc:
                texts.append(page.get_text())
            doc.close()
            return "\n".join(texts)
        except ImportError:
            logger.warning("PyMuPDF (fitz) 未安装，尝试使用 pdfplumber")
            try:
                import pdfplumber
                texts = []
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        texts.append(page.extract_text() or "")
                return "\n".join(texts)
            except ImportError:
                logger.error("PDF 解析失败: 请安装 PyMuPDF (pip install PyMuPDF) 或 pdfplumber")
                return None
    
    def _parse_docx(self, file_path: str) -> Optional[str]:
        """解析 Word 文件"""
        try:
            from docx import Document
            doc = Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # 也提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        paragraphs.append(" | ".join(row_text))
            return "\n".join(paragraphs)
        except ImportError:
            logger.error("Word 解析失败: 请安装 python-docx (pip install python-docx)")
            return None
    
    def _parse_text(self, file_path: str) -> str:
        """解析纯文本文件"""
        encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        # 兜底：用 latin-1（不会抛异常）
        with open(file_path, 'r', encoding='latin-1') as f:
            return f.read()
    
    def _clean_text(self, text: str) -> str:
        """清理文本"""
        # 替换多种空白字符为单个换行
        text = re.sub(r'[\r\n]+', '\n', text)
        # 替换连续空格为单个空格
        text = re.sub(r'[ \t]+', ' ', text)
        # 去掉空行
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join(line for line in lines if line)
        return text.strip()
    
    def _chunk_text(self, text: str) -> List[str]:
        """将文本分块，优先按段落/句子边界切分"""
        if len(text) <= self.chunk_size:
            return [text] if text else []
        
        chunks = []
        # 先按段落分割
        paragraphs = text.split('\n')
        current_chunk = ""
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            
            # 如果单个段落超过 chunk_size，按句子切分
            if len(para) > self.chunk_size:
                # 先保存当前 chunk
                if current_chunk:
                    chunks.append(current_chunk.strip())
                    current_chunk = ""
                
                # 按句子切分大段落
                sentences = re.split(r'(?<=[。！？.!?])\s*', para)
                temp = ""
                for sent in sentences:
                    if len(temp) + len(sent) > self.chunk_size:
                        if temp:
                            chunks.append(temp.strip())
                        temp = sent
                    else:
                        temp += sent
                if temp:
                    chunks.append(temp.strip())
                continue
            
            # 正常段落
            if len(current_chunk) + len(para) + 1 > self.chunk_size:
                chunks.append(current_chunk.strip())
                current_chunk = para
            else:
                current_chunk = (current_chunk + "\n" + para).strip() if current_chunk else para
        
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        # 添加重叠
        if self.chunk_overlap > 0 and len(chunks) > 1:
            overlapped = []
            for i, chunk in enumerate(chunks):
                if i > 0:
                    prev_end = chunks[i - 1][-self.chunk_overlap:]
                    chunk = prev_end + chunk
                overlapped.append(chunk)
            chunks = overlapped
        
        return chunks
    
    @classmethod
    def is_supported(cls, file_path: str) -> bool:
        """检查文件是否支持解析"""
        return Path(file_path).suffix.lower() in cls.SUPPORTED_EXTENSIONS
