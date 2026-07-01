"""
Office 文档导入导出模块单元测试
Export Module Unit Tests
"""

import pytest
import os
import tempfile
import shutil

from medai.export.document_exporter import (
    PaperExporter, GrantProposalExporter,
    ResponseLetterExporter, ProtocolExporter,
)
from medai.export.spreadsheet_exporter import (
    MetaAnalysisExporter, BudgetExporter,
    JournalDatabaseExporter, SurvivalDataExporter,
)
from medai.export.presentation_exporter import (
    ResearchPresentationExporter, ImagingTeachingExporter,
    BioinformaticsReportExporter,
)
from medai.export.document_importer import (
    WordImporter, ExcelImporter,
)


class TestDocumentExporter:
    """Word 文档导出测试"""

    @pytest.fixture
    def temp_dir(self):
        td = tempfile.mkdtemp()
        yield td
        shutil.rmtree(td, ignore_errors=True)

    def test_paper_exporter(self, temp_dir):
        exporter = PaperExporter()
        paper_data = {
            'title': 'Test Paper Title',
            'authors': '张三, 李四',
            'abstract': 'This is a test abstract.',
            'keywords': ['test', 'medical', 'AI'],
            'introduction': 'Introduction section content.',
            'methods': 'Methods section content.',
            'results': 'Results section content.',
            'discussion': 'Discussion section content.',
            'conclusion': 'Conclusion section content.',
            'references': ['Ref 1', 'Ref 2'],
        }
        path = os.path.join(temp_dir, 'test_paper.docx')
        result = exporter.export_paper(paper_data, path)
        assert os.path.exists(result)
        assert os.path.getsize(result) > 0

    def test_grant_proposal_exporter(self, temp_dir):
        exporter = GrantProposalExporter()
        proposal_data = {
            'title': 'Test Grant Proposal',
            'applicant': '张三',
            'institution': '某某医院',
            'rationale': 'Background and rationale content.',
            'research_content': 'Research content details.',
            'objectives': '1. Obj 1\n2. Obj 2',
            'key_problems': 'Key scientific problems to solve.',
            'methodology': 'Method 1\nMethod 2',
            'feasibility': 'Feasibility analysis content.',
            'innovation': 'Innovation points.',
            'budget': {
                'items': [
                    {'name': '设备费', 'amount': 200000, 'notes': '购买仪器'},
                    {'name': '材料费', 'amount': 100000, 'notes': '实验耗材'},
                    {'name': '劳务费', 'amount': 150000, 'notes': '研究生劳务'},
                ],
                'total': 450000,
            },
            'timeline': 'Year 1: Task 1\nYear 2: Task 2',
            'expected_outcomes': '论文3篇\n专利1项',
        }
        path = os.path.join(temp_dir, 'test_proposal.docx')
        result = exporter.export_proposal(proposal_data, path)
        assert os.path.exists(result)
        assert os.path.getsize(result) > 0

    def test_response_letter_exporter(self, temp_dir):
        exporter = ResponseLetterExporter()
        letter_data = {
            'manuscript_id': 'MED-2024-001',
            'title': 'Response to Reviewers',
            'authors': '张三, 李四',
            'responses': [
                {
                    'comment': 'Sample size is too small.',
                    'response': 'We have increased the sample size to 200.',
                    'changes': 'Updated methods section.',
                },
                {
                    'comment': 'Please add more references.',
                    'response': 'We have added 10 more references.',
                    'changes': 'Added references in introduction.',
                },
            ],
        }
        path = os.path.join(temp_dir, 'test_letter.docx')
        result = exporter.export_response_letter(letter_data, path)
        assert os.path.exists(result)
        assert os.path.getsize(result) > 0

    def test_protocol_exporter(self, temp_dir):
        exporter = ProtocolExporter()
        protocol_data = {
            'study_info': {
                'title': 'Test RCT Protocol',
                'phase': 'Phase III',
                'indication': 'Hypertension',
            },
            'study_objectives': {
                'primary': 'Reduce blood pressure',
                'secondary': ['Improve QoL', 'Reduce side effects'],
            },
            'endpoints': {
                'primary': 'SBP change at 12 weeks',
                'secondary': ['DBP change', 'Adverse events'],
            },
            'inclusion_criteria': ['Age 18-75', 'Diagnosed hypertension'],
            'exclusion_criteria': ['Pregnancy', 'Severe organ dysfunction'],
            'randomization': {
                'method': 'Computer-generated',
                'ratio': '1:1',
            },
            'statistical_analysis': {
                'primary_method': 'ANOCOVA',
                'significance_level': 0.05,
            },
        }
        path = os.path.join(temp_dir, 'test_protocol.docx')
        result = exporter.export_protocol(protocol_data, path)
        assert os.path.exists(result)
        assert os.path.getsize(result) > 0


class TestSpreadsheetExporter:
    """Excel 导出测试"""

    @pytest.fixture
    def temp_dir(self):
        td = tempfile.mkdtemp()
        yield td
        shutil.rmtree(td, ignore_errors=True)

    def test_meta_analysis_exporter(self, temp_dir):
        exporter = MetaAnalysisExporter()
        result_data = {
            'studies': [
                {'name': 'Study A', 'effect_size': 0.5, 'se': 0.1, 'weight': 30},
                {'name': 'Study B', 'effect_size': 0.7, 'se': 0.15, 'weight': 25},
                {'name': 'Study C', 'effect_size': 0.6, 'se': 0.12, 'weight': 45},
            ],
            'pooled_effect': 0.58,
            'ci_lower': 0.45,
            'ci_upper': 0.71,
            'heterogeneity': {'i_squared': 25, 'p_value': 0.3},
        }
        path = os.path.join(temp_dir, 'test_meta.xlsx')
        result = exporter.export_meta_analysis(result_data, path)
        assert os.path.exists(result)
        assert os.path.getsize(result) > 0

    def test_budget_exporter(self, temp_dir):
        exporter = BudgetExporter()
        budget_data = {
            'items': [
                {'name': '设备费', 'amount': 200000, 'notes': '仪器'},
                {'name': '材料费', 'amount': 100000, 'notes': '耗材'},
                {'name': '劳务费', 'amount': 150000, 'notes': '人员'},
            ],
            'total': 450000,
        }
        path = os.path.join(temp_dir, 'test_budget.xlsx')
        result = exporter.export_budget(budget_data, path)
        assert os.path.exists(result)
        assert os.path.getsize(result) > 0

    def test_journal_database_exporter(self, temp_dir):
        exporter = JournalDatabaseExporter()
        journals = [
            {'name': 'Nature Medicine', 'if_2023': 87.241, 'category': 'Medicine', 'country': 'UK'},
            {'name': 'Lancet', 'if_2023': 98.4, 'category': 'Medicine', 'country': 'UK'},
            {'name': '中华医学杂志', 'if_2023': 1.5, 'category': 'Medicine', 'country': 'China'},
        ]
        path = os.path.join(temp_dir, 'test_journals.xlsx')
        result = exporter.export_journals(journals, path)
        assert os.path.exists(result)
        assert os.path.getsize(result) > 0

    def test_survival_data_exporter(self, temp_dir):
        exporter = SurvivalDataExporter()
        records = [
            {'patient_id': 'P1', 'time': 12, 'event': 1, 'group': 'A'},
            {'patient_id': 'P2', 'time': 24, 'event': 0, 'group': 'A'},
            {'patient_id': 'P3', 'time': 8, 'event': 1, 'group': 'B'},
        ]
        path = os.path.join(temp_dir, 'test_survival.xlsx')
        result = exporter.export_survival_data(records, path)
        assert os.path.exists(result)
        assert os.path.getsize(result) > 0


class TestPresentationExporter:
    """PPT 导出测试"""

    @pytest.fixture
    def temp_dir(self):
        td = tempfile.mkdtemp()
        yield td
        shutil.rmtree(td, ignore_errors=True)

    def test_research_presentation(self, temp_dir):
        exporter = ResearchPresentationExporter()
        data = {
            'title': 'Research Presentation',
            'subtitle': 'A Study on Medical AI',
            'background': ['Background point 1', 'Background point 2'],
            'methods': ['Method A', 'Method B'],
            'results': ['Result 1', 'Result 2'],
            'discussion': ['Discussion point 1', 'Discussion point 2'],
            'conclusions': ['Conclusion 1'],
            'acknowledgments': ['Acknowledgment to the team'],
        }
        path = os.path.join(temp_dir, 'test_research.pptx')
        result = exporter.export_research_report(data, path)
        assert os.path.exists(result)
        assert os.path.getsize(result) > 0

    def test_imaging_teaching(self, temp_dir):
        exporter = ImagingTeachingExporter()
        signs = [
            {
                'name': 'Ground-Glass Opacity',
                'description': 'Hazy increase in lung attenuation',
                'modalities': ['CT'],
                'anatomy': ['Lung'],
                'diseases': ['COVID-19', 'Pneumonia', 'Pulmonary edema'],
                'severity': 'Variable',
            },
        ]
        path = os.path.join(temp_dir, 'test_teaching.pptx')
        result = exporter.export_teaching(signs, path)
        assert os.path.exists(result)
        assert os.path.getsize(result) > 0

    def test_bioinformatics_report(self, temp_dir):
        exporter = BioinformaticsReportExporter()
        data = {
            'title': 'Genomic Analysis Report',
            'subtitle': 'WES Analysis',
            'sample_info': ['Sample: T-001', 'Type: Tumor'],
            'mutation_summary': ['TP53 p.R273H', 'KRAS p.G12D'],
            'pathways': ['TP53 signaling', 'MAPK cascade'],
            'survival': ['OS median: 24 months'],
        }
        path = os.path.join(temp_dir, 'test_bioinfo.pptx')
        result = exporter.export_bioinformatics_report(data, path)
        assert os.path.exists(result)
        assert os.path.getsize(result) > 0


class TestDocumentImporter:
    """文档导入测试"""

    @pytest.fixture
    def temp_dir(self):
        td = tempfile.mkdtemp()
        yield td
        shutil.rmtree(td, ignore_errors=True)

    def test_word_importer_extract_text(self, temp_dir):
        """测试从Word文档提取文本"""
        from docx import Document as DocxDocument
        path = os.path.join(temp_dir, 'test_doc.docx')
        doc = DocxDocument()
        doc.add_heading('Test Title', level=1)
        doc.add_paragraph('This is a test paragraph.')
        doc.add_paragraph('Second paragraph with more text.')
        doc.save(path)

        importer = WordImporter()
        text = importer.extract_text(path)
        assert 'Test Title' in text
        assert 'test paragraph' in text

    def test_word_importer_structure(self, temp_dir):
        """测试提取文档结构"""
        from docx import Document as DocxDocument
        path = os.path.join(temp_dir, 'test_struct.docx')
        doc = DocxDocument()
        doc.add_heading('Introduction', level=1)
        doc.add_paragraph('Intro text.')
        doc.add_heading('Methods', level=1)
        doc.add_paragraph('Methods text.')
        doc.add_heading('Results', level=1)
        doc.add_paragraph('Results text.')
        doc.save(path)

        importer = WordImporter()
        structure = importer.extract_structure(path)
        assert 'headings' in structure
        assert len(structure['headings']) >= 3

    def test_excel_importer_read_sheet(self, temp_dir):
        """测试读取Excel工作表"""
        from openpyxl import Workbook
        path = os.path.join(temp_dir, 'test_data.xlsx')
        wb = Workbook()
        ws = wb.active
        ws['A1'] = 'Name'
        ws['B1'] = 'Age'
        ws['C1'] = 'Score'
        ws.append(['Alice', 25, 85])
        ws.append(['Bob', 30, 90])
        ws.append(['Charlie', 35, 78])
        wb.save(path)

        importer = ExcelImporter()
        data = importer.read_sheet(path)
        assert len(data) == 3
        assert data[0]['Name'] == 'Alice'
        assert data[1]['Age'] == 30
        assert data[2]['Score'] == 78

    def test_excel_importer_read_all_sheets(self, temp_dir):
        """测试读取所有工作表"""
        from openpyxl import Workbook
        path = os.path.join(temp_dir, 'test_multi.xlsx')
        wb = Workbook()
        ws1 = wb.active
        ws1.title = 'Sheet1'
        ws1['A1'] = 'Col1'
        ws1.append([1, 2])
        ws2 = wb.create_sheet('Sheet2')
        ws2['A1'] = 'ColA'
        ws2.append(['a', 'b'])
        wb.save(path)

        importer = ExcelImporter()
        data = importer.read_all_sheets(path)
        assert 'Sheet1' in data
        assert 'Sheet2' in data
        assert len(data['Sheet1']) == 1
        assert len(data['Sheet2']) == 1

    def test_excel_importer_survival_data(self, temp_dir):
        """测试读取生存数据"""
        from openpyxl import Workbook
        path = os.path.join(temp_dir, 'test_surv.xlsx')
        wb = Workbook()
        ws = wb.active
        ws['A1'] = 'patient_id'
        ws['B1'] = 'time'
        ws['C1'] = 'event'
        ws['D1'] = 'group'
        ws.append(['P1', 12, 1, 'A'])
        ws.append(['P2', 24, 0, 'A'])
        ws.append(['P3', 8, 1, 'B'])
        wb.save(path)

        importer = ExcelImporter()
        data = importer.read_survival_data(path)
        assert len(data) == 3
        assert data[0]['patient_id'] == 'P1'
        assert data[0]['time'] == 12
        assert data[0]['event'] == 1

    def test_excel_importer_journal_database(self, temp_dir):
        """测试读取期刊数据库"""
        from openpyxl import Workbook
        path = os.path.join(temp_dir, 'test_journals.xlsx')
        wb = Workbook()
        ws = wb.active
        ws['A1'] = 'Journal Name'
        ws['B1'] = 'Impact Factor'
        ws['C1'] = 'JCR Quartile'
        ws['D1'] = 'CAS Quartile'
        ws['E1'] = 'Field'
        ws['F1'] = 'OA Policy'
        ws['G1'] = 'Review Period'
        ws.append(['Nature Medicine', 87.241, 'Q1', 'Q1', 'Medicine', 'Hybrid', '2 months'])
        ws.append(['Lancet', 98.4, 'Q1', 'Q1', 'Medicine', 'Subscription', '4 weeks'])
        wb.save(path)

        importer = ExcelImporter()
        data = importer.read_journal_database(path)
        assert len(data) == 2
        assert data[0]['name'] == 'Nature Medicine'
        assert data[1]['impact_factor'] == 98.4
