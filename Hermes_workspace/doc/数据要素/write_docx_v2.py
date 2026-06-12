#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""写入申报书内容到附件1-数据大赛-申报书.docx - FIXED: 从后往前插入"""

import os
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

WORK_DIR = r'D:\workspace\Hermes_workspace\doc\数据要素'
TEMPLATE = os.path.join(WORK_DIR, '附件1-数据大赛-申报书.docx')
OUTPUT = os.path.join(WORK_DIR, '附件1-数据大赛-申报书_已填充.docx')

def read_md(filename):
    path = os.path.join(WORK_DIR, filename)
    with open(path, 'r', encoding='utf-8') as f:
        return f.read()

def extract_section(text, marker):
    idx = text.find(marker)
    if idx == -1:
        return None
    rest = text[idx:].split('\n', 1)
    return rest[1].strip() if len(rest) > 1 else ''

def extract_between(text, start_marker, end_marker=None):
    start = text.find(start_marker)
    if start == -1:
        return None
    s = start + len(start_marker)
    if end_marker:
        e = text.find(end_marker, s)
        return text[s:e].strip() if e != -1 else text[s:].strip()
    return text[s:].strip()

def insert_paragraph_after(paragraph, text=''):
    """Insert a new paragraph after the given paragraph using XML."""
    new_p_elem = OxmlElement('w:p')
    paragraph._p.addnext(new_p_elem)
    from docx.text.paragraph import Paragraph
    new_para = type(paragraph)(new_p_elem, paragraph._parent)
    if text:
        run = new_para.add_run(text)
        run.font.size = Pt(11)
        run.font.name = '宋体'
    return new_para

# Read all files
part1 = read_md('申报书填充内容_part1.md')
part2a = read_md('申报书填充内容_part2a.md')
part2b = read_md('申报书填充内容_part2b.md')
part2cd = read_md('申报书填充内容_part2cd.md')
part3 = read_md('申报书填充内容_part3.md')
part4_security = read_md('申报书填充内容_part4_security.md')

# Extract sections
sections = {}
sections['项目背景'] = extract_section(part1, '### （一）项目背景（限500字）')
sections['应用场景'] = extract_section(part1, '### （二）应用场景（限500字）')
sections['核心优势'] = extract_section(part1, '### （三）核心优势（限1000字）')
sections['数据要素基础'] = extract_section(part2a, '### （一）数据要素基础（限3000字）')
sections['技术路线正文'] = extract_section(part2b, '### （二）技术路线（限4000字）')
sections['数据治理'] = extract_section(part2cd, '### （三）数据治理（限3000字）')
sections['机制创新'] = extract_section(part2cd, '### （四）机制创新与模式创新（限3000字）')
sections['安全保障'] = extract_section(part4_security, '### （五）安全保障（限1000字）')
sections['应用成效'] = extract_between(part3, '## 三、应用成效（限5000字）', '## 四、商业模式（限5000字）')
sections['商业模式'] = extract_between(part3, '## 四、商业模式（限5000字）', None)

# Verify
for key, val in sections.items():
    if val:
        print('OK:', key, '-', len(val), 'chars')
    else:
        print('ERROR: no content for', key)

# Open template
doc = Document(TEMPLATE)

# Build heading position map
h_keys = [
    '（一）项目背景（限500字）',
    '（二）应用场景（限500字）',
    '（三）核心优势（限1000字）',
    '（一）数据要素基础（限3000字）',
    '（二）技术路线（限4000字）',
    '（三）数据治理（限3000字）',
    '（四）机制创新与模式创新（限3000字）',
    '（五）安全保障（限1000字）',
    '三、应用成效（限5000字）',
    '四、商业模式（限5000字）',
]

# Map heading text to section key
h_to_section = {
    '（一）项目背景（限500字）': '项目背景',
    '（二）应用场景（限500字）': '应用场景',
    '（三）核心优势（限1000字）': '核心优势',
    '（一）数据要素基础（限3000字）': '数据要素基础',
    '（二）技术路线（限4000字）': '技术路线正文',
    '（三）数据治理（限3000字）': '数据治理',
    '（四）机制创新与模式创新（限3000字）': '机制创新',
    '（五）安全保障（限1000字）': '安全保障',
    '三、应用成效（限5000字）': '应用成效',
    '四、商业模式（限5000字）': '商业模式',
}

# Find all heading indices in the document
h_indices = []
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    for hk in h_keys:
        if hk in t and ('Heading' in p.style.name):
            h_indices.append((i, hk, p))
            break

# Sort by index (ascending)
h_indices.sort(key=lambda x: x[0])

print()
print('=== 标题位置(原始) ===')
for idx, hk, p in h_indices:
    print(f'  Para {idx}: {hk}')

print()
print('=== 开始写入 (从后往前) ===')

# Process in REVERSE order to avoid index shifting
for start_idx, hk, head_para in reversed(h_indices):
    skey = h_to_section[hk]
    content = sections.get(skey)
    if not content:
        print(f'WARN: no content for {skey}')
        continue
    
    # Insert content line by line AFTER the heading
    lines = content.split('\n')
    ref_para = head_para
    count = 0
    for line in lines:
        stripped = line.strip()
        new_p = insert_paragraph_after(ref_para, stripped)
        count += 1
        ref_para = new_p
    
    print(f'OK {skey}: inserted {count} paragraphs after para {start_idx}')

# Save
doc.save(OUTPUT)
size_kb = os.path.getsize(OUTPUT) / 1024
print()
print(f'OK saved: {OUTPUT}')
print(f'File size: {size_kb:.1f} KB')
