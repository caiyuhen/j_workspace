from docx import Document
import os

os.chdir(r'D:\workspace\Hermes_workspace\doc\数据要素')

path = '附件1-数据大赛-申报书.docx'
doc = Document(path)

for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        print(f'P{i}|{p.style.name}|{p.text}')

print('\n\n========== TABLES ==========')
for ti, table in enumerate(doc.tables):
    print(f'\nTable {ti}: {len(table.rows)} rows x {len(table.columns)} cols')
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            text = cell.text.strip()
            if text:
                print(f'  T{ti}R{ri}C{ci}: {text}')
