#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Final verification of docx content"""

from docx import Document

OUTPUT = r'D:\workspace\Hermes_workspace\doc\数据要素\附件1-数据大赛-申报书_已填充.docx'
doc = Document(OUTPUT)

# Check last 20 paragraphs
print('=== 最后20个段落 ===')
for i in range(max(0, len(doc.paragraphs)-20), len(doc.paragraphs)):
    t = doc.paragraphs[i].text.strip()[:150]
    if t:
        print(f'  Para {i}: {t}')

print()
print('=== 检查关键内容片段 ===')
full_text = '\n'.join(p.text for p in doc.paragraphs)

checks = [
    ('项目背景结尾', '推动中医药行业数据标准化、流通化、智能化转型'),
    ('核心优势', '三层架构体系'),
    ('数据要素基础', '中药成分数据库'),
    ('技术路线', '量子计算'),
    ('数据治理', '五位一体'),
    ('机制创新', '数智辨证画像'),
    ('安全保障', '数据销毁'),
    ('需求痛点', '数据资源"散、乱、孤"'),
    ('质效提升', '辨证准确率'),
    ('社会经济', '直接经济效益'),
    ('综合成效', '复用系数'),
    ('推广示范', '市场潜力'),
    ('模式可持续', '盈利模式'),
    ('财务预测', '净利润'),
    ('社会价值', '扶贫'),
    ('ESG', 'ESG'),
    ('医保衔接', '按疗效付费'),
    ('SROI', 'SROI'),
]

for name, keyword in checks:
    found = keyword in full_text
    status = 'OK' if found else 'MISSING'
    print(f'  {status}: {name} -> "{keyword}"')

print()
# Count total Chinese characters in content
import re
chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', full_text))
print(f'总中文字数: {chinese_chars}')
print(f'总段落数: {len(doc.paragraphs)}')
