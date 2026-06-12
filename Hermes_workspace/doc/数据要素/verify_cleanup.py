#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""验证清理后的docx"""

from docx import Document
import os

PATH = r'D:\workspace\Hermes_workspace\doc\数据要素\附件1-数据大赛-申报书_已填充.docx'

doc = Document(PATH)

print(f'文件大小: {os.path.getsize(PATH)/1024:.1f} KB')
print(f'总段落数: {len(doc.paragraphs)}')

empty_count = 0
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if not t:
        empty_count += 1
    # Show headings
    style = p.style.name if p.style else ''
    if 'Heading' in style:
        print(f'  [{style}] {t[:70]}')

print(f'\n空段落数: {empty_count}')
print()

# Check for any remaining bad text
bad_keywords = [
    '技术架构：介绍参赛作品的顶层设计方案',
    '围绕所选赛题方向，介绍参赛项目',
    '简要介绍参赛作品适用的行业范围',
    '从创新性、有效性和可推广性等方面',
    '描述所申报项目方案是否切中',
    '结合本赛道，描述项目方案',
    '围绕解决方案的市场潜力',
    '说明解决方案的市场策略',
]

found_bad = False
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    for kw in bad_keywords:
        if kw in t and 'Heading' not in (p.style.name if p.style else ''):
            print(f'⚠️ 残留 Para {i}: {t[:80]}...')
            found_bad = True

if not found_bad:
    print('✅ 无模板残留文字')

# Check for remaining markdown placeholders
md_headers = [p.text.strip() for p in doc.paragraphs if p.text.strip().startswith('### ')]
if md_headers:
    print(f'⚠️ Markdown残留: {md_headers[:5]}')
else:
    print('✅ 无Markdown残留占位符')