#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""清理docx中的无用字符、空段落、模板残留"""

from docx import Document
from docx.oxml.ns import qn
import re, os, shutil

WORK_DIR = r'D:\workspace\Hermes_workspace\doc\数据要素'
INPUT = os.path.join(WORK_DIR, '附件1-数据大赛-申报书_已填充.docx')
OUTPUT = os.path.join(WORK_DIR, '附件1-数据大赛-申报书_已填充_清理版.docx')
BACKUP = os.path.join(WORK_DIR, '附件1-数据大赛-申报书_已填充_备份.docx')

shutil.copy2(INPUT, BACKUP)
print(f'已备份: {BACKUP}')

doc = Document(INPUT)

# 模板残留指南文字列表（完整匹配前缀）
guide_prefixes = [
    '围绕所选赛题方向，介绍参赛项目的行业背景',
    '简要介绍参赛作品适用的行业范围及应用场景',
    '从创新性、有效性和可推广性等方面，简要介绍参赛作品的技术优势',
    '介绍参赛作品的顶层设计方案、技术架构',
    '不同赛道根据评价标准可有不同侧重',
    '描述所申报项目方案是否切中所在领域重点',
    '结合本赛道，描述项目方案实现的降本、提效、增质等实际效果',
    '项目落地后带来的经济效益和社会效益',
    '围绕解决方案的市场潜力，开展成长性分析',
    '说明解决方案的市场策略，新模式新业态培育情况',
    '参赛项目申报书',
    '一、参赛团队',
    '二、参赛单位',
    '三、同一参赛单位',
    '四、参赛团队若选择',
    '五、参赛团队须遵守',
    '六、上海分赛主办',
    '七、获得晋级',
    '八、参赛单位、团体或个人',
    '九、大赛期间',
    '十、报名参赛',
    '1.知识产权情况',
    '2.合同情况',
    '3.其他证明材料',
]

# 模板残留markdown占位符
md_placeholders = [
    '### （一）项目背景（限500字）',
    '### （二）应用场景（限500字）',
    '### （三）核心优势（限1000字）',
    '### （一）数据要素基础（限3000字）',
    '### （二）技术路线（限4000字）',
    '### （三）数据治理（限3000字）',
    '### （四）机制创新与模式创新（限3000字）',
    '### （五）安全保障（限1000字）',
    '### （一）需求痛点',
    '### （二）质效提升成效',
    '### （三）经济社会效益',
    '### （四）综合成效评估',
    '### （一）推广示范价值',
    '### （二）模式可持续性',
    '### （三）财务预测与投资回报',
    '### （四）社会价值商业模式',
    '**（5）基因表达与生物组学数据**',
    '**（6）中医四诊数据（约50万例）**',
    '**（7）针灸经络数据**',
    '**（8）药品流通与质量控制数据**',
]

# 已知的具体模板说明文字（完整匹配）
known_bad_texts = [
    '描述所申报项目方案是否切中所在领域重点、难点、堵点等重要需求。项目解决问题的重要程度、问题的普遍性/代表性、问题解决程度和影响范围。',
    '项目落地后带来的经济效益和社会效益。',
    '结合本赛道，描述项目方案实现的降本、提效、增质等实际效果。包括但不限于项目如何体现数据要素提质增效、发挥数据赋能价值的情况。',
    '说明解决方案的市场策略，新模式新业态培育情况，包括数据来源、数据要素利用模式、产品价格、成本核算、盈利模式及稳定性、未来应用空间、推广渠道、宣传方式等，如有可提供成本、收入、未来应用空间等测算说明。',
    '围绕解决方案的市场潜力，开展成长性分析。如潜在用户规模、行业领域、市场份额等情况。项目是否形成具有可复制、可推广的运用数据要素赋能行业的解决方案或应用模式。项目是否具备数据治理标准推广水平或数据流通生态构建水平。',
    '介绍参赛作品的顶层设计方案、技术架构等。数据资源赛道阐述数据资源载体和应用系统。数据基础设施赛道着重阐述基础设施的技术架构和部署情况。',
    '本项目构建了覆盖"标准—质量—安全—流通—伦理"五位一体的中医药数据治理体系，确保数据全生命周期的规范化管理与合规使用。',
    '一、项目概述',
    '二、解决方案',
    '三、应用成效',
    '四、商业模式',
    '五、附件',
    '不同赛道根据评价标准可有不同侧重。',
]

def should_remove(para):
    """判断段落是否应该被删除"""
    text = para.text.strip()
    
    # 空段落
    if not text:
        return True
    if re.match(r'^\s+$', text):
        return True
    
    style = para.style.name if para.style else ''
    
    # 保留标题段落
    if 'Heading' in style or 'toc' in style.lower():
        return False
    
    # 模板指南文字
    for prefix in guide_prefixes:
        if text.startswith(prefix):
            return True
    
    # Markdown占位符
    if text in md_placeholders:
        return True
    
    # ###开头的markdown残留
    if text.startswith('### ') or text.startswith('## '):
        return True
    
    # **开头的模板残留（纯占位符行）
    if text.startswith('**（') and text.endswith('**'):
        return True
    
    # 已知模板说明文字
    for bad in known_bad_texts:
        if text == bad:
            return True
    
    # 纯数字/标点行
    if re.match(r'^[\d\s\.\-—·,，。、]+$', text):
        return True
    
    return False

# 收集要删除的段落
remove_indices = set()
for i, p in enumerate(doc.paragraphs):
    if should_remove(p):
        remove_indices.add(i)

print(f'待删除段落数: {len(remove_indices)}')

# 从后往前删除（XML级别操作）
body = doc.element.body
para_elems = body.findall(qn('w:p'))

# 反向排序确保索引不偏移
for idx in sorted(remove_indices, reverse=True):
    if idx < len(para_elems):
        body.remove(para_elems[idx])

# 保存
doc.save(OUTPUT)
print(f'已保存: {OUTPUT}')
print(f'文件大小: {os.path.getsize(OUTPUT)/1024:.1f} KB')

# 验证
doc2 = Document(OUTPUT)
total = len(doc2.paragraphs)
empty = sum(1 for p in doc2.paragraphs if not p.text.strip())
print(f'清理后段落数: {total}')
print(f'清理后空段落数: {empty}')

# 检查所有段落是否都干净
print()
print('=== 剩余疑似问题段落 ===')
for i, p in enumerate(doc2.paragraphs):
    t = p.text.strip()
    if not t:
        continue
    style = p.style.name if p.style else ''
    if any(p in t or t.startswith(p) for p in guide_prefixes):
        print(f'  WARN Para {i}: "{t[:80]}"')
    if any(t == m for m in md_placeholders):
        print(f'  WARN Para {i}: placeholder "{t[:80]}"')
    if t.startswith('### ') or t.startswith('## '):
        print(f'  WARN Para {i}: markdown header "{t[:80]}"')