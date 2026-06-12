#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""深度验证写入的docx内容"""

from docx import Document
import os

OUTPUT = r'D:\workspace\Hermes_workspace\doc\数据要素\附件1-数据大赛-申报书_已填充.docx'

doc = Document(OUTPUT)
print(f'总段落数: {len(doc.paragraphs)}')
print()

# Find all heading-like paragraphs
print('=== 所有标题/章节 ===')
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if 'Heading' in p.style.name or t.startswith('##') or t.startswith('（一）') or t.startswith('（二）') or t.startswith('（三）') or t.startswith('（四）') or t.startswith('（五）') or t.startswith('三、') or t.startswith('四、') or t.startswith('###'):
        if t:
            print(f'  Para {i}: [{p.style.name}] {t[:100]}')

print()
print('=== 最后15个段落 ===')
for i in range(max(0, len(doc.paragraphs)-15), len(doc.paragraphs)):
    t = doc.paragraphs[i].text.strip()
    if t:
        print(f'  Para {i}: {t[:120]}')
