#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""分析docx中的无用字符和空段落"""

from docx import Document
import re

OUTPUT = r'D:\workspace\Hermes_workspace\doc\数据要素\附件1-数据大赛-申报书_已填充.docx'
doc = Document(OUTPUT)

print('=== 空段落/无内容段落统计 ===')
empty_count = 0
placeholder_count = 0
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if not t:
        empty_count += 1
    elif '（' in t and len(t) < 30 and not p.style.name.startswith('Heading'):
        # Likely leftover template placeholder like （一）项目背景（限500字）
        placeholder_count += 1
        print(f'  Empty/Near-empty para {i}: style={p.style.name} text="{t[:60]}"')

print(f'\n总空段落数: {empty_count}')
print(f'疑似残留模板占位符段落: {placeholder_count}')

print()
print('=== 每个章节附近的段落内容 ===')

# Find all headings
headings = []
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if 'Heading' in p.style.name and t:
        headings.append((i, t))

for idx, title in headings:
    print(f'\n[H] Para {idx}: {title[:80]}')
    # Show next 5 paragraphs
    for j in range(idx+1, min(idx+8, len(doc.paragraphs))):
        t = doc.paragraphs[j].text.strip()
        if t:
            print(f'  Para {j}: {t[:120]}...' if len(t) > 120 else f'  Para {j}: {t}')
        else:
            print(f'  Para {j}: [空]')

print()
print('=== 模板残留占位符（不含用户内容的段落） ===')
# Look for template boilerplate - paragraphs that look generic
keywords_to_check = ['参赛', '团队', '报名', '赛区', '赛题', '原则上', '同一参赛']
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if any(kw in t for kw in keywords_to_check):
        print(f'  Para {i}: {t[:100]}')
