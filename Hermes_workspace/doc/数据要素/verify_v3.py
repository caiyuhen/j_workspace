#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""验证docx内容完整性"""
from docx import Document
import re

OUTPUT = r'D:\workspace\Hermes_workspace\doc\数据要素\附件1-数据大赛-申报书_已填充.docx'
doc = Document(OUTPUT)

full_text = '\n'.join(p.text for p in doc.paragraphs)

# Check all sections present
checks = {
    '项目背景': '数据要素×',
    '应用场景': '三大核心应用场景',
    '核心优势': '三层架构',
    '数据要素基础': '中药成分',
    '技术路线': '量子计算',
    '数据治理': '五位一体',
    '机制创新': '数智辨证画像',
    '安全保障': '数据销毁',
    '应用成效': '辨证准确率',
    '商业模式': '盈利模式',
}

print('=== 章节关键词验证 ===')
for name, kw in checks.items():
    found = kw in full_text
    print(f'  {"OK" if found else "FAIL"}: {name} -> "{kw}"')

# Count Chinese chars
chinese = len(re.findall(r'[\u4e00-\u9fff]', full_text))
print(f'\n总中文字数: {chinese}')
print(f'总段落数: {len(doc.paragraphs)}')

# Quick sample: find all headings
print('\n=== 标题结构 ===')
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if 'Heading' in p.style.name:
        print(f'  Para {i}: {t[:100]}')
