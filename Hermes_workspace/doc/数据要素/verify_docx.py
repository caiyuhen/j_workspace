#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""验证写入的docx内容"""

from docx import Document
import os

OUTPUT = r'D:\workspace\Hermes_workspace\doc\数据要素\附件1-数据大赛-申报书_已填充.docx'

doc = Document(OUTPUT)
print(f'总段落数: {len(doc.paragraphs)}')
print(f'文件大小: {os.path.getsize(OUTPUT)/1024:.1f} KB')
print()

# Check key section starts
check_indices = [51, 53, 55, 59, 61, 65, 67, 69, 71, 79]
print('=== 各章节标题 ===')
for i in check_indices:
    if i < len(doc.paragraphs):
        t = doc.paragraphs[i].text.strip()[:120]
        print(f'  Para {i}: {t}')

print()
print('=== 内容抽样验证 ===')
# Check content after each heading
for i, h_idx in enumerate(check_indices):
    next_idx = check_indices[i+1] if i+1 < len(check_indices) else len(doc.paragraphs)
    # Get first content paragraph after heading
    if h_idx + 1 < len(doc.paragraphs):
        t = doc.paragraphs[h_idx + 1].text.strip()[:100]
        if t:
            print(f'  After para {h_idx}: "{t}..."')
    # Get last content paragraph before next heading
    if next_idx > h_idx + 2:
        last_idx = next_idx - 1
        if last_idx < len(doc.paragraphs):
            t = doc.paragraphs[last_idx].text.strip()[:100]
            if t:
                print(f'  Before para {next_idx}: "...{t[-80:]}"')
    print()

# Check total section content completeness by counting non-empty paragraphs
print('=== 非空段落统计 ===')
for i, start_h in enumerate(check_indices):
    end_h = check_indices[i+1] if i+1 < len(check_indices) else len(doc.paragraphs)
    content_paras = []
    total_chars = 0
    for j in range(start_h, end_h):
        t = doc.paragraphs[j].text.strip()
        if t:
            content_paras.append(t)
            total_chars += len(t)
    print(f'  章节{start_h}-{end_h}: {len(content_paras)}个段落, 共{total_chars}字')
