#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""清理残留的1个模板指南段落"""

from docx import Document
from docx.oxml.ns import qn
import os

INPUT = r'D:\workspace\Hermes_workspace\doc\数据要素\附件1-数据大赛-申报书_已填充_清理版.docx'

doc = Document(INPUT)
body = doc.element.body
para_elems = body.findall(qn('w:p'))

# Find the paragraph containing template guide text
for idx, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if '技术架构：介绍参赛作品的顶层设计方案' in t:
        print(f'Found at paragraph {idx}: "{t[:60]}..."')
        # Get the XML element at the same logical position
        # Use the actual paragraph element
        elem = para_elems[idx]
        body.remove(elem)
        print('Removed!')
        break

out = INPUT  # overwrite
doc.save(out)
print(f'Saved: {out}')

# Verify
doc2 = Document(INPUT)
still_bad = [p.text.strip() for p in doc2.paragraphs if '技术架构：介绍' in p.text]
print(f'Remaining bad paragraphs: {len(still_bad)}')
print(f'Total paragraphs: {len(doc2.paragraphs)}')