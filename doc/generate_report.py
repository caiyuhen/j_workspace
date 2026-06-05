from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 创建新文档
doc = Document()

# 设置默认字体
doc.styles['Normal'].font.name = '宋体'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 添加标题
title = doc.add_paragraph()
title_run = title.add_run('常规监查访视报告')
title_run.font.size = Pt(16)
title_run.font.bold = True
title_run.font.name = '黑体'
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 创建基本信息表格
table1 = doc.add_table(rows=8, cols=2)
table1.style = 'Table Grid'

# 基本信息数据
basic_info = [
    ('项目编号/Project No.', 'CT-2024-015'),
    ('项目名称/Title', '评价注射用XX药物急性缺血性卒中治疗的III期临床试验'),
    ('申办单位/Sponsor', 'XX制药有限公司'),
    ('中心编号/名称：', '北京天坛医院（中心编号：008）'),
    ('中心负责人：', '王XX 主任医师'),
    ('访视起止日期：', '2026年4月21日'),
    ('报告首次递交日期', '2026年4月21日'),
    ('：', ''),
]

for i, (key, value) in enumerate(basic_info):
    row = table1.rows[i]
    row.cells[0].text = key
    row.cells[1].text = value
    row.cells[0].width = Cm(5)
    row.cells[1].width = Cm(10)

# 添加空行
doc.add_paragraph()

# 参加访视人员表格
doc.add_paragraph('参加访视人员').bold = True
table2 = doc.add_table(rows=3, cols=2)
table2.style = 'Table Grid'
table2.rows[0].cells[0].text = '姓名'
table2.rows[0].cells[1].text = '工作职责'
table2.rows[1].cells[0].text = '张XX'
table2.rows[1].cells[1].text = 'CRA（监查员）'
table2.rows[2].cells[0].text = '李XX'
table2.rows[2].cells[1].text = 'PM（项目经理，陪同）'

# 本次监查所接触的研究中心参与研究的人员表格
doc.add_paragraph()
doc.add_paragraph('本次监查所接触的研究中心参与研究的人员').bold = True
table3 = doc.add_table(rows=4, cols=2)
table3.style = 'Table Grid'
table3.rows[0].cells[0].text = '姓名'
table3.rows[0].cells[1].text = '工作职责'
table3.rows[1].cells[0].text = '王XX'
table3.rows[1].cells[1].text = '主要研究者（PI）'
table3.rows[2].cells[0].text = '陈XX'
table3.rows[2].cells[1].text = '研究医生'
table3.rows[3].cells[0].text = '刘XX'
table3.rows[3].cells[1].text = '研究护士/协调员'

# 受试者入组情况
doc.add_page_break()
doc.add_paragraph('受试者入组情况').bold = True
table4 = doc.add_table(rows=2, cols=6)
table4.style = 'Table Grid'
table4.rows[0].cells[0].text = '计划入组数'
table4.rows[0].cells[1].text = '筛选例数'
table4.rows[0].cells[2].text = '入组例数'
table4.rows[0].cells[3].text = '进行中例数'
table4.rows[0].cells[4].text = '完成例数'
table4.rows[0].cells[5].text = '脱落例数'
table4.rows[1].cells[0].text = '60'
table4.rows[1].cells[1].text = '48'
table4.rows[1].cells[2].text = '45'
table4.rows[1].cells[3].text = '18'
table4.rows[1].cells[4].text = '25'
table4.rows[1].cells[5].text = '2'

# 脱落病例描述
doc.add_paragraph()
doc.add_paragraph('脱落病例描述').bold = True
table5 = doc.add_table(rows=3, cols=2)
table5.style = 'Table Grid'
table5.rows[0].cells[0].text = '受试者编号'
table5.rows[0].cells[1].text = '原因/备注'
table5.rows[1].cells[0].text = '008-012'
table5.rows[1].cells[1].text = '受试者自行退出（个人原因）'
table5.rows[2].cells[0].text = '008-023'
table5.rows[2].cells[1].text = '失访（联系电话变更，无法联系）'

# 严重不良事件（SAE）
doc.add_paragraph()
doc.add_paragraph('严重不良事件（SAE）').bold = True
table6 = doc.add_table(rows=2, cols=3)
table6.style = 'Table Grid'
table6.rows[0].cells[0].text = '受试者编号'
table6.rows[0].cells[1].text = '事件'
table6.rows[0].cells[2].text = '是否24小时汇报'
table6.rows[1].cells[0].text = '008-035'
table6.rows[1].cells[1].text = '轻度颅内出血（已恢复，与试验药物无关）'
table6.rows[1].cells[2].text = '是'

# 监查内容
doc.add_page_break()
doc.add_paragraph('监查内容').bold = True

# 1、研究进度
doc.add_paragraph('1、研究进度').bold = True
table7 = doc.add_table(rows=4, cols=4)
table7.style = 'Table Grid'
table7.rows[0].cells[0].text = ''
table7.rows[0].cells[1].text = '是'
table7.rows[0].cells[2].text = '否'
table7.rows[0].cells[3].text = 'N/A'

table7.rows[1].cells[0].text = '1）入组是否与预期相符合？'
table7.rows[1].cells[1].text = '☑'
table7.rows[1].cells[2].text = '☐'
table7.rows[1].cells[3].text = '☐'

table7.rows[2].cells[0].text = '2）和研究人员一起评估了受试者招募过程中的问题？'
table7.rows[2].cells[1].text = '☑'
table7.rows[2].cells[2].text = '☐'
table7.rows[2].cells[3].text = '☐'

table7.rows[3].cells[0].text = '3）研究人员及时更新了受试者筛选/入选表？'
table7.rows[3].cells[1].text = '☑'
table7.rows[3].cells[2].text = '☐'
table7.rows[3].cells[3].text = '☐'

p = doc.add_paragraph()
p.add_run('备注：').bold = True
p.add_run('目前入组进度良好，已完成75%入组计划。研究团队表示按目前进度预计可在2个月内完成全部入组。')

# 2、知情同意
doc.add_paragraph()
doc.add_paragraph('2、知情同意').bold = True
table8 = doc.add_table(rows=5, cols=4)
table8.style = 'Table Grid'
table8.rows[0].cells[0].text = ''
table8.rows[0].cells[1].text = '是'
table8.rows[0].cells[2].text = '否'
table8.rows[0].cells[3].text = 'N/A'

table8.rows[1].cells[0].text = '1）研究者正在使用的知情同意书已经获得IEC/IRB的批准？'
table8.rows[1].cells[1].text = '☑'
table8.rows[1].cells[2].text = '☐'
table8.rows[1].cells[3].text = '☐'

table8.rows[2].cells[0].text = '2）每位受试者在开始试验之前均按照法规要求签署了知情同意书？'
table8.rows[2].cells[1].text = '☑'
table8.rows[2].cells[2].text = '☐'
table8.rows[2].cells[3].text = '☐'

table8.rows[3].cells[0].text = '3）知情同意整个过程是否正确，是否被及时和正确记录？'
table8.rows[3].cells[1].text = '☑'
table8.rows[3].cells[2].text = '☐'
table8.rows[3].cells[3].text = '☐'

table8.rows[4].cells[0].text = '4）若知情同意书为纸质版，那么研究者保留了完整知情同意书，同时受试者获取副本？'
table8.rows[4].cells[1].text = '☑'
table8.rows[4].cells[2].text = '☐'
table8.rows[4].cells[3].text = '☐'

p = doc.add_paragraph()
p.add_run('备注：').bold = True
p.add_run('本中心使用V2.0版知情同意书（2024年3月15日伦理批准）。已核查受试者008-028至008-045的知情同意书签署情况，签署完整规范。')

# ISF版本信息
doc.add_paragraph()
doc.add_paragraph('ISF版本').bold = True
table_isf = doc.add_table(rows=4, cols=2)
table_isf.style = 'Table Grid'
table_isf.rows[0].cells[0].text = '版本时间'
table_isf.rows[0].cells[1].text = '2024年3月20日'
table_isf.rows[1].cells[0].text = '伦理批件时间'
table_isf.rows[1].cells[1].text = '2024年3月15日'
table_isf.rows[2].cells[0].text = '受试者编号'
table_isf.rows[2].cells[1].text = '知情时间'
table_isf.rows[3].cells[0].text = '008-035'
table_isf.rows[3].cells[1].text = '2026年3月28日'

# 3、方案及法规的依从性
doc.add_page_break()
doc.add_paragraph('3、方案及法规的依从性').bold = True
table9 = doc.add_table(rows=4, cols=4)
table9.style = 'Table Grid'
table9.rows[0].cells[0].text = ''
table9.rows[0].cells[1].text = '是'
table9.rows[0].cells[2].text = '否'
table9.rows[0].cells[3].text = 'N/A'

table9.rows[1].cells[0].text = '1）研究者依从试验方案（包括最新版本）？'
table9.rows[1].cells[1].text = '☑'
table9.rows[1].cells[2].text = '☐'
table9.rows[1].cells[3].text = '☐'

table9.rows[2].cells[0].text = '2）研究者依从相关政策法规？'
table9.rows[2].cells[1].text = '☑'
table9.rows[2].cells[2].text = '☐'
table9.rows[2].cells[3].text = '☐'

table9.rows[3].cells[0].text = '3）研究者在向IEC/IRB或相关药监部门报告履行了其职能？'
table9.rows[3].cells[1].text = '☑'
table9.rows[3].cells[2].text = '☐'
table9.rows[3].cells[3].text = '☐'

p = doc.add_paragraph()
p.add_run('备注：').bold = True
p.add_run('研究团队对方案依从性良好，本次访视未发现重大方案偏离。')

# 4、研究中心的人员和设施
doc.add_paragraph()
doc.add_paragraph('4、研究中心的人员和设施').bold = True
table10 = doc.add_table(rows=6, cols=4)
table10.style = 'Table Grid'
table10.rows[0].cells[0].text = ''
table10.rows[0].cells[1].text = '是'
table10.rows[0].cells[2].text = '否'
table10.rows[0].cells[3].text = 'N/A'

table10.rows[1].cells[0].text = '1）研究人员和设施充足？'
table10.rows[1].cells[1].text = '☑'
table10.rows[1].cells[2].text = '☐'
table10.rows[1].cells[3].text = '☐'

table10.rows[2].cells[0].text = '2）试验所需设备得到及时补充？'
table10.rows[2].cells[1].text = '☑'
table10.rows[2].cells[2].text = '☐'
table10.rows[2].cells[3].text = '☐'

table10.rows[3].cells[0].text = '3）研究人员发生变化？'
table10.rows[3].cells[1].text = '☐'
table10.rows[3].cells[2].text = '☑'
table10.rows[3].cells[3].text = '☐'

table10.rows[4].cells[0].text = '4）实验室样本的采集、处理、标记、储存和运送过程及数据记录符合相关要求？'
table10.rows[4].cells[1].text = '☑'
table10.rows[4].cells[2].text = '☐'
table10.rows[4].cells[3].text = '☐'

table10.rows[5].cells[0].text = '5）实验室正常值范围、单位、检测方法、实验室室间质评证书是否有变化？'
table10.rows[5].cells[1].text = '☐'
table10.rows[5].cells[2].text = '☐'
table10.rows[5].cells[3].text = '☑'

p = doc.add_paragraph()
p.add_run('备注：').bold = True
p.add_run('研究团队人员稳定，未发生变更。中心设施设备齐全，药品储存条件符合要求。')

# 5、安全性监查
doc.add_page_break()
doc.add_paragraph('5、安全性监查').bold = True
table11 = doc.add_table(rows=9, cols=4)
table11.style = 'Table Grid'
table11.rows[0].cells[0].text = ''
table11.rows[0].cells[1].text = '是'
table11.rows[0].cells[2].text = '否'
table11.rows[0].cells[3].text = 'N/A'

table11.rows[1].cells[0].text = '1）发现包括受试者安全性方面的任何问题并与研究人员进行了讨论？'
table11.rows[1].cells[1].text = '☑'
table11.rows[1].cells[2].text = '☐'
table11.rows[1].cells[3].text = '☐'

table11.rows[2].cells[0].text = '2）对以上问题与研究人员就解决办法和/或预防措施达成共识？'
table11.rows[2].cells[1].text = '☑'
table11.rows[2].cells[2].text = '☐'
table11.rows[2].cells[3].text = '☐'

table11.rows[3].cells[0].text = '3）是否自从上次监查以来有SAE发生？'
table11.rows[3].cells[1].text = '☑'
table11.rows[3].cells[2].text = '☐'
table11.rows[3].cells[3].text = '☐'

table11.rows[4].cells[0].text = '4）这些SAE首次报告是否在24小时内上报？'
table11.rows[4].cells[1].text = '☑'
table11.rows[4].cells[2].text = '☐'
table11.rows[4].cells[3].text = '☐'

table11.rows[5].cells[0].text = '5）是否有提交SAE的随访、总结报告？'
table11.rows[5].cells[1].text = '☑'
table11.rows[5].cells[2].text = '☐'
table11.rows[5].cells[3].text = '☐'

table11.rows[6].cells[0].text = '6）SAE报告上的源数据是否有被核查、核对？'
table11.rows[6].cells[1].text = '☑'
table11.rows[6].cells[2].text = '☐'
table11.rows[6].cells[3].text = '☐'

table11.rows[7].cells[0].text = '7）是否有任何仍待执行的对于SAE的后续处理？'
table11.rows[7].cells[1].text = '☐'
table11.rows[7].cells[2].text = '☑'
table11.rows[7].cells[3].text = '☐'

table11.rows[8].cells[0].text = '8）近期获得了新的安全性信息？'
table11.rows[8].cells[1].text = '☐'
table11.rows[8].cells[2].text = '☐'
table11.rows[8].cells[3].text = '☑'

p = doc.add_paragraph()
p.add_run('备注：').bold = True
p.add_run('本次访视期间共发现1例SAE（008-035受试者轻度颅内出血），已确认研究者按照GCP要求在24小时内完成上报。SAE报告与源数据一致，随访报告完整。')

# 6、实验用药及其他研究相关资料
doc.add_paragraph()
doc.add_paragraph('6、实验用药及其他研究相关资料').bold = True
table12 = doc.add_table(rows=7, cols=4)
table12.style = 'Table Grid'
table12.rows[0].cells[0].text = ''
table12.rows[0].cells[1].text = '是'
table12.rows[0].cells[2].text = '否'
table12.rows[0].cells[3].text = 'N/A'

table12.rows[1].cells[0].text = '1）药品发放记录准确且被及时更新？'
table12.rows[1].cells[1].text = '☑'
table12.rows[1].cells[2].text = '☐'
table12.rows[1].cells[3].text = '☐'

table12.rows[2].cells[0].text = '2）试验用药按照方案规定保存？'
table12.rows[2].cells[1].text = '☑'
table12.rows[2].cells[2].text = '☐'
table12.rows[2].cells[3].text = '☐'

table12.rows[3].cells[0].text = '3）有破盲并且相关信息得以正确记录？'
table12.rows[3].cells[1].text = '☐'
table12.rows[3].cells[2].text = '☐'
table12.rows[3].cells[3].text = '☑'

table12.rows[4].cells[0].text = '4）IB、ICF、Protocol有更新版本？'
table12.rows[4].cells[1].text = '☐'
table12.rows[4].cells[2].text = '☑'
table12.rows[4].cells[3].text = '☐'

table12.rows[5].cells[0].text = '7）试验用药及其他相关研究资料是否供应充足？'
table12.rows[5].cells[1].text = '☑'
table12.rows[5].cells[2].text = '☐'
table12.rows[5].cells[3].text = '☐'

table12.rows[6].cells[0].text = '8）本次访视是否有试验用药品及其他相关研究资料的回收报告或者销毁报告？'
table12.rows[6].cells[1].text = '☐'
table12.rows[6].cells[2].text = '☑'
table12.rows[6].cells[3].text = '☐'

p = doc.add_paragraph()
p.add_run('备注：').bold = True
p.add_run('试验用药储存于专用药柜，温度记录完整（2-8℃）。当前库存：试验药85支，对照药72支。发放记录准确，与CRF一致。IB、ICF、Protocol为当前有效版本，无更新。')

# 7、CRF核查
doc.add_page_break()
doc.add_paragraph('7、CRF核查').bold = True
table13 = doc.add_table(rows=5, cols=4)
table13.style = 'Table Grid'
table13.rows[0].cells[0].text = ''
table13.rows[0].cells[1].text = '是'
table13.rows[0].cells[2].text = '否'
table13.rows[0].cells[3].text = 'N/A'

table13.rows[1].cells[0].text = '1）原始数据/源文件是否都存在？'
table13.rows[1].cells[1].text = '☑'
table13.rows[1].cells[2].text = '☐'
table13.rows[1].cells[3].text = '☐'

table13.rows[2].cells[0].text = '2）是否CRF完成率、疑问表的解决情况符合研究要求？'
table13.rows[2].cells[1].text = '☑'
table13.rows[2].cells[2].text = '☐'
table13.rows[2].cells[3].text = '☐'

table13.rows[3].cells[0].text = '3）化验单和检查报告单是否被及时接收、审阅并存档？'
table13.rows[3].cells[1].text = '☑'
table13.rows[3].cells[2].text = '☐'
table13.rows[3].cells[3].text = '☐'

table13.rows[4].cells[0].text = '4) 是否进行了SDR / SDV？'
table13.rows[4].cells[1].text = '☑'
table13.rows[4].cells[2].text = '☐'
table13.rows[4].cells[3].text = '☐'

p = doc.add_paragraph()
p.add_run('备注：').bold = True
p.add_run('本次访视对受试者008-028至008-045进行了100% SDV核查，共计18例受试者。发现3处数据录入错误，已现场纠正。')

# 7、ISF核查
doc.add_paragraph()
doc.add_paragraph('7、ISF核查').bold = True
table14 = doc.add_table(rows=4, cols=4)
table14.style = 'Table Grid'
table14.rows[0].cells[0].text = ''
table14.rows[0].cells[1].text = '是'
table14.rows[0].cells[2].text = '否'
table14.rows[0].cells[3].text = 'N/A'

table14.rows[1].cells[0].text = '1）是否有收集任何研究相关的文件？'
table14.rows[1].cells[1].text = '☑'
table14.rows[1].cells[2].text = '☐'
table14.rows[1].cells[3].text = '☐'

table14.rows[2].cells[0].text = '2）是否有按照本中心伦理委员会要求完成文件递交？'
table14.rows[2].cells[1].text = '☑'
table14.rows[2].cells[2].text = '☐'
table14.rows[2].cells[3].text = '☐'

table14.rows[3].cells[0].text = '3）有更新研究者文件夹？'
table14.rows[3].cells[1].text = '☑'
table14.rows[3].cells[2].text = '☐'
table14.rows[3].cells[3].text = '☐'

p = doc.add_paragraph()
p.add_run('备注：').bold = True
p.add_run('ISF文件齐全，已补充2024年度伦理递交记录。')

# 原始数据/源文件问题汇总
doc.add_page_break()
doc.add_paragraph('原始数据/源文件问题汇总').bold = True
table15 = doc.add_table(rows=4, cols=3)
table15.style = 'Table Grid'
table15.rows[0].cells[0].text = '受试者编号'
table15.rows[0].cells[1].text = '访视周期'
table15.rows[0].cells[2].text = '问题描述'

table15.rows[1].cells[0].text = '008-032'
table15.rows[1].cells[1].text = '访视4'
table15.rows[1].cells[2].text = 'CRF中血压值录入错误（源数据138/85mmHg，CRF录入为128/85mmHg），已更正'

table15.rows[2].cells[0].text = '008-038'
table15.rows[2].cells[1].text = '访视2'
table15.rows[2].cells[2].text = 'CRF中用药日期记录不完整，已补充完整'

table15.rows[3].cells[0].text = '008-041'
table15.rows[3].cells[1].text = '筛选期'
table15.rows[3].cells[2].text = '病史记录中合并用药遗漏记录，已补充'

# 补充备注及其它问题
doc.add_paragraph()
doc.add_paragraph('补充备注及其它问题').bold = True

p = doc.add_paragraph()
p.add_run('1、试验用药品的管理（补充、分发、回收、计算、保存和处理等）\n')
p.add_run('   药品管理规范，库存清点准确。建议下次访视前进行一次药品回收。')

p = doc.add_paragraph()
p.add_run('2、试验中的问题与难点或严重依从性偏离（违反伦理或影响试验安全性和有效性评价）\n')
p.add_run('   无严重方案偏离。发现2例轻微偏离（受试者008-029、008-037访视窗口期超时3天），均已记录并说明原因，不影响受试者安全性。')

p = doc.add_paragraph()
p.add_run('3、与研究人员的其它相关讨论或提供的信息\n')
p.add_run('   与研究团队讨论了后续受试者招募计划，研究团队表示有信心在预计时间内完成入组。')

p = doc.add_paragraph()
p.add_run('4、其它重大发现或数据质疑\n')
p.add_run('   无。')

# 相关措施与跟踪
doc.add_paragraph()
doc.add_paragraph('相关措施与跟踪').bold = True
table16 = doc.add_table(rows=4, cols=3)
table16.style = 'Table Grid'
table16.rows[0].cells[0].text = '需采取的措施'
table16.rows[0].cells[1].text = '实施负责人'
table16.rows[0].cells[2].text = '完成期限'

table16.rows[1].cells[0].text = '补充受试者008-041的合并用药记录'
table16.rows[1].cells[1].text = '陈XX医生'
table16.rows[1].cells[2].text = '2026年4月25日'

table16.rows[2].cells[0].text = '更新研究者文件夹中的年度伦理递交文件'
table16.rows[2].cells[1].text = '刘XX护士'
table16.rows[2].cells[2].text = '2026年4月28日'

table16.rows[3].cells[0].text = '安排下次监查访视'
table16.rows[3].cells[1].text = '张XX（CRA）'
table16.rows[3].cells[2].text = '2026年5月30日'

p = doc.add_paragraph()
p.add_run('前次监查遗留的相关问题是否解决？ ☑是，☐否')

# 保存文档
output_path = r'D:\doc\JD-FOR-SMV-03常规监查访视报告_天坛医院_20260421.docx'
doc.save(output_path)
print(f"文档已生成：{output_path}")
