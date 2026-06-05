#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成美化的 Word 简历文档
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

# 设置文件路径
INPUT_FILE = r'D:\doc\蔡宇衡的简历-AI 技术总监优化版.md'
OUTPUT_FILE = r'D:\doc\蔡宇衡的简历-AI 技术总监优化版.docx'

def set_font(run, font_name='微软雅黑', size=Pt(12), bold=False, color=None):
    """设置字体样式"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = size
    run.font.bold = bold
    if color:
        run.font.color.rgb = color

def add_heading(doc, text, level=1):
    """添加标题"""
    if level == 1:
        heading = doc.add_heading(text, level=1)
        run = heading.runs[0]
        set_font(run, size=Pt(16), bold=True, color=RGBColor(46, 117, 182))
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(12)
    elif level == 2:
        heading = doc.add_heading(text, level=2)
        run = heading.runs[0]
        set_font(run, size=Pt(14), bold=True, color=RGBColor(31, 78, 121))
        heading.paragraph_format.space_before = Pt(10)
        heading.paragraph_format.space_after = Pt(8)
    elif level == 3:
        heading = doc.add_heading(text, level=3)
        run = heading.runs[0]
        set_font(run, size=Pt(13), bold=True, color=RGBColor(91, 155, 213))
        heading.paragraph_format.space_before = Pt(8)
        heading.paragraph_format.space_after = Pt(6)
    return heading

def add_bullet_list(doc, items):
    """添加项目符号列表"""
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        set_font(run, size=Pt(11))
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)

def add_table(doc, data, has_header=True):
    """添加表格"""
    table = doc.add_table(rows=len(data), cols=2)
    table.style = 'Table Grid'
    
    for i, row_data in enumerate(data):
        row = table.rows[i]
        cells = row.cells
        
        for j, text in enumerate(row_data):
            cell = cells[j]
            cell.width = Inches(3.5)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            run = paragraph.add_run(text)
            if has_header and i == 0:
                set_font(run, size=Pt(11), bold=True, color=RGBColor(31, 78, 121))
                # 简化表格背景设置
                from docx.oxml import OxmlElement
                tc = cell._tc
                tcPr = OxmlElement('w:tcPr')
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), 'D5E8F0')
                tcPr.append(shd)
                tc.insert(0, tcPr)
            else:
                set_font(run, size=Pt(11))
    
    doc.add_paragraph()

def main():
    # 创建文档
    doc = Document()
    
    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
    
    # ========== 标题部分 ==========
    # 姓名
    name_p = doc.add_paragraph()
    name_run = name_p.add_run('蔡宇衡')
    set_font(name_run, size=Pt(22), bold=True, color=RGBColor(46, 117, 182))
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 联系方式
    contact_p = doc.add_paragraph()
    contact_run = contact_p.add_run('📞 13810357924  |  📧 caiyuheng81@outlook.com  |  📍 北京')
    set_font(contact_run, size=Pt(12), color=RGBColor(46, 46, 46))
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # ========== 求职意向 ==========
    add_heading(doc, '求职意向', level=1)
    add_table(doc, [
        ['目标职位', 'AI 技术总监 / 技术 VP / CTO'],
        ['期望薪资', '60-85k×16 薪'],
        ['工作地点', '北京'],
        ['行业方向', '人工智能、医疗健康、SaaS 平台']
    ])
    
    # ========== 核心优势 ==========
    add_heading(doc, '核心优势', level=1)
    
    # 技术战略与架构规划
    add_heading(doc, '🎯 技术战略与架构规划', level=3)
    add_bullet_list(doc, [
        '22 年软件开发经验，15 年技术管理背景，擅长制定中长期技术演进路线图，主导前端框架升级、微服务架构改造、API 网关解耦、BFF 层建设等系统性技术决策',
        '具备从 0 到 1 组建百人级产研团队经验，建立跨部门研发管理体系，对系统连续性、AI 落地、研发交付节奏负总责',
        '精通微服务架构设计，能在关键技术问题上做兜底决策，管理外包运维伙伴服务质量与 SLA'
    ])
    
    # 大模型应用与技术落地
    add_heading(doc, '🤖 大模型应用与技术落地', level=3)
    add_bullet_list(doc, [
        '深度实践 LLM 应用开发，精通 Prompt Engineering、RAG 检索增强、Agent 编排、Function Calling、Skill 框架设计，独立交付过 AI Agent 系统上线',
        '熟悉 LangChain、LlamaIndex、Dify、AutoGen 等主流 LLM 应用框架，具备多模态医学影像大模型优化与训练策略设计经验',
        '主导构建企业级 RAG 知识库系统，实现文档分词、向量存储、语义检索的全链路技术闭环'
    ])
    
    # 医疗行业与 SaaS 经验
    add_heading(doc, '🏥 医疗行业与 SaaS 经验', level=3)
    add_bullet_list(doc, [
        '15 年深耕医疗健康领域，熟悉 ICH GCP、FDA 21 CFR Part 11、CDISC/SDTM 等合规标准',
        '主导开发 CTMS+EDC 临床试验管理系统，具备私有部署 + 多租户 SaaS 系统完整架构经验',
        '熟悉企业微信应用生态，有医疗行业政企 SaaS 产品落地经验'
    ])
    
    # 跨组织协作与团队管理
    add_heading(doc, '👥 跨组织协作与团队管理', level=3)
    add_bullet_list(doc, [
        '擅长在涉及多个利益相关方（合作技术单位、机构 IT、业务方）的复杂协作场景下，维持稳定合作关系同时守住团队工作节奏与决策边界',
        '牵头 40+ 真实世界研究项目，主导国家级课题申报及行业标准制定（NMPA/CDE/北京大学合作）'
    ])
    
    doc.add_page_break()
    
    # ========== 工作经历 ==========
    add_heading(doc, '工作经历', level=1)
    
    # 安顿健康科技
    p = doc.add_paragraph()
    run = p.add_run('安顿健康科技有限公司 | 高级算法工程师（技术负责人）')
    set_font(run, bold=True, color=RGBColor(31, 78, 121))
    p = doc.add_paragraph()
    run = p.add_run('2025.07 - 至今')
    set_font(run, size=Pt(11), color=RGBColor(102, 102, 102))
    
    doc.add_paragraph('核心职责：').add_run().font.bold = True
    add_bullet_list(doc, [
        '制定 AI 技术中长期发展规划，主导跨部门技术协作机制建设',
        '负责智能穿戴设备健康预测系统架构设计，管理外包开发团队 SLA 交付质量'
    ])
    
    doc.add_paragraph('重点项目：').add_run().font.bold = True
    add_heading(doc, '1. 中医智诊机器人系统（AI Agent 实践）', level=3)
    add_bullet_list(doc, [
        '技术架构决策：基于 LangChain + RAG + AutoGen 构建多 Agent 协作框架，设计 Skill 框架扩展机制',
        'RAG 知识库构建：基于《黄帝内经》《伤寒杂病论》等中医典籍建立知识体系（8 万 + 方剂、2000+ 证候、9000+ 中药材）',
        '多模态融合：结合图像识别（望诊）、传感器数据（切诊）、智能问答（问诊）实现四诊法数字化',
        '成果：完成系统从架构设计到上线的全流程，支持临床辅助决策'
    ])
    
    # 圣方上海
    p = doc.add_paragraph()
    run = p.add_run('圣方 (上海) 医药研发有限公司 | 高级数据科学家 / 高级算法工程师（技术负责人）')
    set_font(run, bold=True, color=RGBColor(31, 78, 121))
    p = doc.add_paragraph()
    run = p.add_run('2021.05 - 2025.04')
    set_font(run, size=Pt(11), color=RGBColor(102, 102, 102))
    
    doc.add_paragraph('核心职责：').add_run().font.bold = True
    add_bullet_list(doc, [
        '制定 AI 技术演进路线图，主导跨部门研发管理体系建设',
        '负责大模型应用方向技术决策，推动算法创新到产品落地全流程贯通'
    ])
    
    doc.add_paragraph('关键技术成果：').add_run().font.bold = True
    add_bullet_list(doc, [
        '技术战略：主导 AI 技术中长期规划，建立从算法研发到产品注册交付的完整闭环',
        '大模型平台：优化医学影像大模型架构，推动多模态融合研究，提升模型临床泛化能力',
        '商业化落地：主导 10+ 款 AI 医疗产品从原型设计到 NMPA 注册交付的全周期管理',
        '科研生态：牵头与国内外机构合作研究，参与制定 3 项行业标准',
        '团队建设：搭建高绩效研发团队，建立技术人才选拔培养机制'
    ])
    
    # 心医国际
    p = doc.add_paragraph()
    run = p.add_run('心医国际 | 副总裁 / 技术 VP')
    set_font(run, bold=True, color=RGBColor(31, 78, 121))
    p = doc.add_paragraph()
    run = p.add_run('2020.09 - 2021.05')
    set_font(run, size=Pt(11), color=RGBColor(102, 102, 102))
    
    add_bullet_list(doc, [
        '主导跨部门技术团队管理，构建全链路研发体系',
        'SaaS 平台架构：完成脑科学及肿瘤专科 SaaS 平台架构设计，支持私有部署 + 多租户模式',
        'CDSS 系统：构建基于知识图谱的临床决策支持系统及 AI 康复系统',
        '联邦学习平台：开发数据交易平台，实现隐私保护下的多中心协作',
        '区域医疗协同：陕西/贵州省级项目支持 100+ 医院急救体系协同数据交互',
        '产品落地：推动 3 款核心产品在医疗机构快速落地，药物警戒系统 300+ 医院分布式部署'
    ])
    
    # 北京首佑医学科技
    p = doc.add_paragraph()
    run = p.add_run('北京首佑医学科技 | 大数据人工智能总监')
    set_font(run, bold=True, color=RGBColor(31, 78, 121))
    p = doc.add_paragraph()
    run = p.add_run('2018.09 - 2020.09')
    set_font(run, size=Pt(11), color=RGBColor(102, 102, 102))
    
    add_bullet_list(doc, [
        '主导 AI 技术团队建设与技术路线规划',
        '知识图谱：通过 NLP 技术解析 150 万 + 电子病历构建精神疾病知识图谱',
        '核心算法：主导语音识别（CNN/RNN）与 MRI 影像分析模型研发，支撑国自然课题',
        '产品落地：设计双抗药物浓度动态监测算法，完成 6 医院临床落地',
        '行业标准：AI 辅助诊断框架技术方案入选 3 个省级医疗标准，诊断准确率提升 95%'
    ])
    
    # 安华亿能医疗影像科技
    p = doc.add_paragraph()
    run = p.add_run('安华亿能医疗影像科技 | CTO')
    set_font(run, bold=True, color=RGBColor(31, 78, 121))
    p = doc.add_paragraph()
    run = p.add_run('2012.04 - 2018.09')
    set_font(run, size=Pt(11), color=RGBColor(102, 102, 102))
    
    add_bullet_list(doc, [
        '制定公司技术战略，管理 40+ 人研发团队',
        '技术创新：首创全球颈动脉三维超声诊断系统，斑块自动检测准确率 98.2%',
        '多模态平台：开发 CT/MRI/超声跨模态融合技术，搭建 DICOM 智能分析系统',
        'SaaS 平台：开发云诊所影像 SaaS 平台，完成 130+ 家三甲医院 PACS 系统无缝对接',
        '医疗器械：主导 II 类医疗器械从概念到 NMPA 认证全流程，建立设计控制闭环'
    ])
    
    doc.add_page_break()
    
    # ========== 核心项目经验 ==========
    add_heading(doc, '核心项目经验', level=1)
    
    # AI Agent 系统
    p = doc.add_paragraph()
    run = p.add_run('AI Agent 系统（药物警戒智能体） | 主导架构设计  2024.01 - 2024.11')
    set_font(run, bold=True, color=RGBColor(31, 78, 121))
    
    doc.add_paragraph('项目背景：构建基于大模型的药物警戒系统，实现 AE/SAE 自动采集、分级、多语言翻译、全球分发')
    doc.add_paragraph('技术架构决策：').add_run().font.bold = True
    add_bullet_list(doc, [
        'Agent 框架：LangChain（工具集成）+ AutoGen（多 Agent 协作）+ Dify（工作流编排）',
        'RAG 检索：基于 LlamaIndex 实现定向检索，Zep 向量存储 + MemGPT 记忆逻辑控制',
        '决策框架：多 LLM 支持（GPT/Claude/Gemini/DeepSeek/Doubao）+ PDDL 规划算法',
        '提示词工程：PromptPerfect 自动优化 + BERT 语义理解偏差修正'
    ])
    
    doc.add_paragraph('核心成果：').add_run().font.bold = True
    add_bullet_list(doc, [
        '完成药物 AE/SAE 自动采集与分级，支持多语言翻译与全球分公司分发',
        '研究数据医学稽查 Agent 定制稽查规则，自动生成稽查报告',
        '✅ 系统已上线运行，实现从数据采集到报告生成的全流程自动化'
    ])
    
    # CTMS+EDC 系统
    p = doc.add_paragraph()
    run = p.add_run('CTMS+EDC 临床试验管理系统 | 架构师  进行中')
    set_font(run, bold=True, color=RGBColor(31, 78, 121))
    
    doc.add_paragraph('项目背景：主导设计临床试验管理系统，符合 ICH GCP、FDA 21 CFR Part 11、CDISC/SDTM 标准')
    doc.add_paragraph('技术架构：').add_run().font.bold = True
    add_bullet_list(doc, [
        '后端：Node.js/TypeScript + Express.js + Prisma ORM + PostgreSQL',
        '前端：React + Vite + Ant Design + Zustand',
        '核心模块：EdcTemplate、CrfForm、CrfData、AdverseEvent、角色权限、工时管理、项目收支、全流程审批',
        '数据合规：6 类 CSV/JSON 导出功能（CDISC 标准 eCRF 与 SDTM 导出）',
        '系统特性：私有部署 + 多租户架构，支持企业微信集成'
    ])
    
    doc.add_paragraph('当前进展：').add_run().font.bold = True
    add_bullet_list(doc, [
        '后端 30+ 模块、前端 21 页面已完成',
        '正在推进 RAG 检索系统构建（文档分词 + 向量库导入）',
        '调研 CDISC/SDTM 标准合规实现，设计 Code Dictionary 映射机制'
    ])
    
    doc.add_page_break()
    
    # ========== 技术能力矩阵 ==========
    add_heading(doc, '技术能力矩阵', level=1)
    add_table(doc, [
        ['大模型应用', 'Prompt Engineering、RAG、Agent 编排、Function Calling、Skill 框架、LangChain、LlamaIndex、Dify、AutoGen、Zep、MemGPT'],
        ['后端架构', 'Node.js/TypeScript、Java、Python、Go、微服务架构、API 网关、BFF 层、Express.js、Spring Boot'],
        ['前端技术', 'React、Vue、前端框架升级、Vite、Ant Design、Zustand'],
        ['数据架构', 'PostgreSQL、MySQL、MongoDB、Redis、Kafka、Spark、Hadoop、数据中台'],
        ['AI/算法', 'Transformer、PyTorch、TensorFlow、RNN、CNN、知识图谱、联邦学习、多模态融合'],
        ['云原生', 'Docker、Kubernetes、Jenkins、HPA、边缘计算 K3S、云原生存储'],
        ['医疗合规', 'ICH GCP E6(R2)、FDA 21 CFR Part 11、GDPR、CDISC、SDTM、OMOP、CDMP'],
        ['企业应用', '企业微信 API、微信生态、政企 SaaS、私有部署、多租户系统']
    ])
    
    doc.add_page_break()
    
    # ========== 教育背景 ==========
    add_heading(doc, '教育背景', level=1)
    add_bullet_list(doc, [
        '山东大学 | 硕士 · 数据科学与大数据技术 | 2004.07 - 2007.06',
        '北京科技大学 | 本科 · 计算机应用 | 1998.07 - 2002.07'
    ])
    
    # ========== 资格证书 ==========
    add_heading(doc, '资格证书', level=1)
    add_bullet_list(doc, [
        '高级人工智能训练师',
        '高级健康管理师',
        '大学英语四级（读写精通）'
    ])
    
    # ========== 附加说明 ==========
    add_heading(doc, '附加说明', level=1)
    add_bullet_list(doc, [
        '系统重构经验：主导过多次老系统重构（前端框架升级、微服务改造、数据中台建设）',
        '外包管理经验：具备外包运维伙伴服务质量与 SLA 管理实战经验',
        '跨组织协作：40+ 项目涉及多方协作（NMPA/CDE/北京大学/三甲医院/药企/保险公司）',
        '技术标准：参与制定 3+ 项行业标准，技术方案入选 3+ 个省级标准',
        '知识产权：获得 3+ 项发明专利，主导国自然课题、"十三五"国家重点研发计划'
    ])
    
    # 保存文档
    doc.save(OUTPUT_FILE)
    print('✅ Word 文档生成成功！')
    print(f'📄 文件路径：{OUTPUT_FILE}')
    print(f'📊 文件大小：{os.path.getsize(OUTPUT_FILE) / 1024:.1f} KB')

if __name__ == '__main__':
    main()
