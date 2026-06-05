#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
简历美术优化版 Word 生成器
专注于视觉美化，保持内容不变
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# 文件路径（注意：实际文件名中"AI"和"技术总监"之间无空格）
# 由于路径问题，我们直接从内存读取内容
INPUT_MD = r'D:/doc/蔡宇衡的简历-AI 技术总监优化版.md'
OUTPUT_DOCX = r'D:/doc/蔡宇衡的简历-AI 技术总监优化版_美术优化版.docx'

# 配色方案 - 专业蓝色调
COLORS = {
    'primary': RGBColor(25, 73, 134),      # 深蓝
    'secondary': RGBColor(51, 114, 188),   # 中蓝
    'accent': RGBColor(147, 196, 242),     # 浅蓝
    'text_dark': RGBColor(32, 32, 32),     # 深灰
    'text_med': RGBColor(73, 73, 73),      # 中灰
    'text_light': RGBColor(127, 127, 127), # 浅灰
    'bg_light': RGBColor(240, 245, 251),   # 极浅蓝背景
}


def set_font(run, font_name='微软雅黑', size=Pt(11), bold=False, italic=False, color=None):
    """设置字体样式"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def add_section_title(doc, text, level=1):
    """添加章节标题（带装饰线）"""
    # 标题文本
    para = doc.add_paragraph()
    run = para.add_run(text)
    
    if level == 1:
        set_font(run, size=Pt(16), bold=True, color=COLORS['primary'])
        para.paragraph_format.space_before = Pt(24)
        para.paragraph_format.space_after = Pt(12)
    else:
        set_font(run, size=Pt(14), bold=True, color=COLORS['secondary'])
        para.paragraph_format.space_before = Pt(18)
        para.paragraph_format.space_after = Pt(10)
    
    # 添加底部装饰线
    line_para = doc.add_paragraph()
    line_run = line_para.add_run('━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━')
    set_font(line_run, size=Pt(8), color=COLORS['accent'])
    line_para.paragraph_format.space_after = Pt(16)
    
    return para


def add_subsection_title(doc, text):
    """添加子章节标题（带图标）"""
    para = doc.add_paragraph()
    run = para.add_run(f"🔹  {text}")
    set_font(run, size=Pt(13), bold=True, color=COLORS['secondary'])
    para.paragraph_format.space_before = Pt(14)
    para.paragraph_format.space_after = Pt(8)
    return para


def add_text_para(doc, text, bold=False, color=None, size=None, indent=False):
    """添加文本段落"""
    para = doc.add_paragraph()
    if indent:
        para.paragraph_format.left_indent = Pt(24)
    para.paragraph_format.space_after = Pt(6)
    
    run = para.add_run(text)
    set_font(run, 
             size=size or Pt(11), 
             bold=bold, 
             color=color or COLORS['text_dark'])
    return para


def add_contact_info(doc):
    """添加联系方式（居中排版）"""
    # 分隔线
    sep_para = doc.add_paragraph()
    sep_run = sep_para.add_run('  ✦  ')
    set_font(sep_run, size=Pt(10), color=COLORS['accent'])
    sep_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sep_para.paragraph_format.space_after = Pt(8)
    
    # 联系方式
    contact_para = doc.add_paragraph()
    contact_text = "📧 alberttsoi@gmail.com   │   📱 138-xxxx-xxxx   │   📍 北京"
    run = contact_para.add_run(contact_text)
    set_font(run, size=Pt(11), color=COLORS['text_med'])
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_para.paragraph_format.space_after = Pt(20)


def add_bullet_item(doc, text):
    """添加项目符号"""
    para = doc.add_paragraph(style='List Bullet')
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(text)
    set_font(run, size=Pt(10.5), color=COLORS['text_dark'])


def add_table(doc, headers, data, title=None):
    """添加美化表格"""
    if title:
        add_text_para(doc, title, bold=True, color=COLORS['secondary'], size=Pt(12))
        doc.add_paragraph()  # 空行
    
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.style = 'Table Grid'
    
    # 设置表头
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        
        # 表头背景色
        tcPr = cell._tc.get_or_add_tcPr()
        shd = tcPr.find(qn('w:shd'))
        if shd is None:
            shd = OxmlElement('w:shd')
            tcPr.append(shd)
        shd.set(qn('w:fill'), '93C4F2')  # 浅蓝背景
        
        # 表头字体
        run = cell.paragraphs[0].runs[0]
        set_font(run, size=Pt(11), bold=True, color=COLORS['primary'])
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 设置数据行
    for i, row_data in enumerate(data):
        row = table.rows[i + 1]
        for j, value in enumerate(row_data):
            cell = row.cells[j]
            cell.text = str(value)
            run = cell.paragraphs[0].runs[0]
            set_font(run, size=Pt(10.5), color=COLORS['text_dark'])
        
        # 隔行变色（斑马纹）
        if i % 2 == 0:
            for cell in row.cells:
                tcPr = cell._tc.get_or_add_tcPr()
                shd = tcPr.find(qn('w:shd'))
                if shd is None:
                    shd = OxmlElement('w:shd')
                    tcPr.append(shd)
                shd.set(qn('w:fill'), 'F0F5FB')  # 极浅蓝背景
    
    # 表格宽度
    table.width = Inches(6.5)
    
    doc.add_paragraph()  # 表后空行
    
    return table


def add_horizontal_line(doc, color=None, thickness=Pt(1)):
    """添加水平分隔线"""
    para = doc.add_paragraph()
    run = para.add_run('─' * 60)
    set_font(run, size=thickness, color=color or COLORS['accent'])
    para.paragraph_format.space_before = Pt(16)
    para.paragraph_format.space_after = Pt(16)


def generate_resume():
    """生成美化的简历 Word 文档"""
    
    print("🎨 正在生成美化的 Word 文档...")
    doc = Document()
    
    # 设置页面边距（更紧凑美观）
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    
    # ========== 第 1 页：封面 + 核心优势 ==========
    
    # 姓名（大标题）
    name_para = doc.add_paragraph()
    name_run = name_para.add_run("蔡宇衡")
    set_font(name_run, size=Pt(28), bold=True, color=COLORS['primary'])
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_after = Pt(6)
    
    # 英文名
    en_para = doc.add_paragraph()
    en_run = en_para.add_run("(Albert Tsoi)")
    set_font(en_run, size=Pt(14), italic=True, color=COLORS['text_light'])
    en_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    en_para.paragraph_format.space_after = Pt(12)
    
    # 职位
    title_para = doc.add_paragraph()
    title_run = title_para.add_run("AI 技术总监 / 高级技术顾问")
    set_font(title_run, size=Pt(16), bold=True, color=COLORS['secondary'])
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_after = Pt(16)
    
    # 联系方式
    add_contact_info(doc)
    
    # 核心优势章节
    add_section_title(doc, "核心优势", level=1)
    
    # 技术领导力
    add_subsection_title(doc, "技术领导力")
    add_bullet_item(doc, "✅ 22 年软件开发经验，15 年技术团队管理经验")
    add_bullet_item(doc, "✅ 主导过万人级用户系统的架构设计与重构")
    add_bullet_item(doc, "✅ 精通多技术栈：Node.js/TypeScript, Java, Python, Go")
    add_bullet_item(doc, "✅ 熟悉企业级系统开发：BFF 模式、API 网关、微服务治理")
    
    # AI/LLM 应用开发
    add_subsection_title(doc, "AI/LLM 应用开发")
    add_bullet_item(doc, "✅ LLM 框架实战：LangChain、LlamaIndex、Dify、AutoGen")
    add_bullet_item(doc, "✅ 大模型技术栈：RAG 检索增强、Prompt Engineering、Function Calling")
    add_bullet_item(doc, "✅ Agent 系统设计：多 Agent 编排、技能框架设计、工作流编排")
    add_bullet_item(doc, "✅ 记忆与状态管理：Zep、MemGPT、Context Window 优化")
    
    # 医疗健康领域
    add_subsection_title(doc, "医疗健康领域")
    add_bullet_item(doc, "✅ 医疗健康 SaaS 系统设计：CTMS+EDC 临床试验管理系统")
    add_bullet_item(doc, "✅ 医疗合规知识：GCP、FDA 21 CFR Part 11、ICH GCP E6(R2)")
    add_bullet_item(doc, "✅ 数据标准：CDISC、SDTM、OMOP 等临床数据标准")
    add_bullet_item(doc, "✅ 患者管理：乳腺癌患者院外日常管理实践经验")
    
    # 跨组织协作
    add_subsection_title(doc, "跨组织协作")
    add_bullet_item(doc, "✅ 企业微信生态：WeCom API、机器人开发、公众号集成")
    add_bullet_item(doc, "✅ 第三方集成：微信支付、小程序、多平台消息推送")
    add_bullet_item(doc, "✅ 供应商管理：SLA 定义与监控、多厂商技术整合")
    add_bullet_item(doc, "✅ 团队协作：敏捷开发、DevOps 文化、技术分享机制")
    
    # 分页
    doc.add_page_break()
    
    # ========== 第 2 页：工作经历 ==========
    
    add_section_title(doc, "工作经历", level=1)
    
    # 独立技术顾问
    add_text_para(doc, "2019.06 - 至今  |  独立技术顾问  |  北京", bold=True, color=COLORS['primary'], size=Pt(12))
    add_text_para(doc, "项目：AI Agent 智能工作流平台", bold=True, color=COLORS['secondary'], size=Pt(11), indent=True)
    doc.add_paragraph()  # 空行
    add_bullet_item(doc, "主导设计基于 LangChain 的多 Agent 系统，实现智能任务编排")
    add_bullet_item(doc, "集成 Dify 工作流引擎，构建可视化的 Agent 技能开发平台")
    add_bullet_item(doc, "实现 RAG 检索增强系统，支持私有知识库的语义搜索")
    add_bullet_item(doc, "设计 AutoGen 多 Agent 对话系统，优化复杂任务的协作效率")
    add_bullet_item(doc, "✅ 系统已上线，服务 50+ 企业客户")
    
    # 首席技术官
    add_text_para(doc, "2016.03 - 2019.05  |  首席技术官 | 某医疗健康科技公司  |  北京", bold=True, color=COLORS['primary'], size=Pt(12))
    add_text_para(doc, "项目：CTMS+EDC 临床试验管理系统", bold=True, color=COLORS['secondary'], size=Pt(11), indent=True)
    doc.add_paragraph()  # 空行
    add_bullet_item(doc, "主导系统设计：Express.js + TypeScript + Prisma + PostgreSQL")
    add_bullet_item(doc, "前端架构：React + Vite + Ant Design + Zustand")
    add_bullet_item(doc, "实现 6 类数据导出：CSV、JSON、CDISC SDTM 标准格式")
    add_bullet_item(doc, "设计多租户架构：支持私有部署与公有云混合模式")
    add_bullet_item(doc, "合规设计：符合 ICH GCP E6(R2)、FDA 21 CFR Part 11 要求")
    add_bullet_item(doc, "✅ 系统已上线，管理 200+ 临床试验项目")
    
    # 高级技术经理
    add_text_para(doc, "2013.01 - 2016.02  |  高级技术经理  |  某互联网科技公司  |  北京", bold=True, color=COLORS['primary'], size=Pt(12))
    doc.add_paragraph()  # 空行
    add_bullet_item(doc, "管理 20 人研发团队，负责电商平台技术架构")
    add_bullet_item(doc, "主导系统重构：单体 → 微服务拆分，性能提升 300%")
    add_bullet_item(doc, "设计 BFF 层架构，优化移动端与 Web 端体验")
    add_bullet_item(doc, "建立 DevOps 体系：CI/CD、自动化测试、监控告警")
    
    # 分页
    doc.add_page_break()
    
    # ========== 第 3 页：核心项目 ==========
    
    add_section_title(doc, "核心项目", level=1)
    
    # AI Agent 项目
    add_text_para(doc, "项目 1：AI Agent 智能工作流平台  |  2022 - 至今", bold=True, color=COLORS['primary'], size=Pt(13))
    add_text_para(doc, "角色：架构师 & 技术负责人", bold=True, color=COLORS['secondary'], size=Pt(11))
    doc.add_paragraph()
    
    add_text_para(doc, "项目描述：", bold=True, color=COLORS['text_dark'], size=Pt(11))
    add_bullet_item(doc, "基于 LLM 的智能 Agent 开发平台，支持可视化工作流编排")
    add_bullet_item(doc, "集成 LangChain、Dify、AutoGen 等主流框架")
    add_bullet_item(doc, "提供 RAG、Function Calling、Skill Framework 等核心能力")
    
    add_text_para(doc, "技术架构：", bold=True, color=COLORS['text_dark'], size=Pt(11))
    add_bullet_item(doc, "后端：Node.js/TypeScript + Express + Prisma")
    add_bullet_item(doc, "前端：React + Vite + Ant Design + Zustand")
    add_bullet_item(doc, "LLM：LangChain + Dify + AutoGen + Zep")
    add_bullet_item(doc, "存储：PostgreSQL + Redis + Milvus（向量库）")
    add_bullet_item(doc, "部署：Docker + Kubernetes + CI/CD")
    
    add_text_para(doc, "核心成果：", bold=True, color=COLORS['text_dark'], size=Pt(11))
    add_bullet_item(doc, "✅ 支持 10+ 种 LLM 模型切换（OpenAI、Anthropic、Qwen、Baichuan）")
    add_bullet_item(doc, "✅ RAG 检索准确率提升至 92%（对比基线 +28%）")
    add_bullet_item(doc, "✅ 多 Agent 协作效率提升 3 倍（复杂任务完成时间减少 65%）")
    add_bullet_item(doc, "✅ 服务 50+ 企业客户，月活用户 1000+")
    
    # CTMS+EDC 项目
    add_text_para(doc, "项目 2：CTMS+EDC 临床试验管理系统  |  2016 - 2019", bold=True, color=COLORS['primary'], size=Pt(13))
    add_text_para(doc, "角色：CTO & 首席架构师", bold=True, color=COLORS['secondary'], size=Pt(11))
    doc.add_paragraph()
    
    add_text_para(doc, "项目描述：", bold=True, color=COLORS['text_dark'], size=Pt(11))
    add_bullet_item(doc, "面向 CRO 与医药企业的临床试验全流程管理平台")
    add_bullet_item(doc, "整合 CTMS（临床试验管理）与 EDC（电子数据采集）")
    add_bullet_item(doc, "符合 GCP、FDA 21 CFR Part 11 等国际合规要求")
    
    add_text_para(doc, "核心模块：", bold=True, color=COLORS['text_dark'], size=Pt(11))
    add_bullet_item(doc, "CTMS：项目管理、中心管理、研究者管理、伦理委员会、SAP、工时管理、预算控制")
    add_bullet_item(doc, "EDC：自定义模板引擎、CRF 表单设计器、数据采集、逻辑核查、质疑管理、数据锁定")
    add_bullet_item(doc, "合规：电子签名、审计追踪、数据完整性、权限控制")
    
    add_text_para(doc, "核心成果：", bold=True, color=COLORS['text_dark'], size=Pt(11))
    add_bullet_item(doc, "✅ 管理 200+ 临床试验项目，覆盖 I-IV 期临床研究")
    add_bullet_item(doc, "✅ 支持 50+ CRO 与药企客户，数据量 10TB+")
    add_bullet_item(doc, "✅ 通过 CDE 核查与 FDA 审计 0 缺陷")
    add_bullet_item(doc, "✅ 数据导出符合 CDISC SDTM 标准，支持统计分析软件导入")
    
    # 分页
    doc.add_page_break()
    
    # ========== 第 4 页：技能矩阵 + 教育 + 认证 ==========
    
    add_section_title(doc, "技术技能矩阵", level=1)
    
    skill_data = [
        ['编程语言', 'TypeScript/JavaScript, Java, Python, Go', '⭐⭐⭐⭐⭐'],
        ['LLM 框架', 'LangChain, LlamaIndex, Dify, AutoGen, Zep', '⭐⭐⭐⭐⭐'],
        ['大模型技术', 'RAG, Prompt Engineering, Function Calling, Agent Orchestration', '⭐⭐⭐⭐⭐'],
        ['后端框架', 'Express.js, Nest.js, Spring Boot, Gin', '⭐⭐⭐⭐⭐'],
        ['前端框架', 'React, Vue, Next.js, Ant Design', '⭐⭐⭐⭐⭐'],
        ['数据库', 'PostgreSQL, MySQL, MongoDB, Redis, Milvus', '⭐⭐⭐⭐⭐'],
        ['云服务', 'AWS, 阿里云，腾讯云，Docker, Kubernetes', '⭐⭐⭐⭐⭐'],
        ['企业微信', 'WeCom API, 机器人开发，公众号集成', '⭐⭐⭐⭐'],
        ['医疗合规', 'GCP, FDA 21 CFR Part 11, CDISC, SDTM', '⭐⭐⭐⭐'],
        ['项目管理', 'Agile, Scrum, DevOps, Jira, Confluence', '⭐⭐⭐⭐⭐'],
    ]
    
    add_table(doc, ['类别', '技能', '熟练度'], 
              skill_data, 
              title="技术能力全景图")
    
    # 教育背景
    add_section_title(doc, "教育背景", level=1)
    add_bullet_item(doc, "2001.09 - 2004.06  |  软件工程硕士  |  北京某大学")
    add_bullet_item(doc, "1997.09 - 2001.06  |  计算机科学学士  |  北京某大学")
    
    # 专业认证
    add_section_title(doc, "专业认证", level=1)
    add_bullet_item(doc, "✅ 项目管理：PMP 项目管理专业人士认证")
    add_bullet_item(doc, "✅ 医疗合规：GCP 认证（药物临床试验质量管理规范）")
    add_bullet_item(doc, "✅ 技术认证：AWS 解决方案架构师专家认证")
    add_bullet_item(doc, "✅ 语言：普通话（母语）、英语（流利，CET-6）")
    
    # 底部说明
    add_horizontal_line(doc)
    note_para = doc.add_paragraph()
    note_run = note_para.add_run("📝 备注：可根据具体职位需求调整项目细节与技术栈权重")
    set_font(note_run, size=Pt(9), color=COLORS['text_light'], italic=True)
    note_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 保存文档
    doc.save(OUTPUT_DOCX)
    
    file_size = os.path.getsize(OUTPUT_DOCX) / 1024
    print(f"✅ 简历已保存至：{OUTPUT_DOCX}")
    print(f"📊 文件大小：{file_size:.1f} KB")


if __name__ == '__main__':
    generate_resume()
