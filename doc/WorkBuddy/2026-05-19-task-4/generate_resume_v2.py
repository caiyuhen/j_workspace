#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
简历 Word 生成器 v2 - 专业美化版
基于 Markdown 内容生成美化的 Word 简历文档
"""

import os
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docxcompose.composer import Composer

# 配置
INPUT_MD = r'D:\doc\蔡宇衡的简历-AI 技术总监优化版.md'
OUTPUT_DOCX = r'D:\doc\蔡宇衡的简历-AI 技术总监优化版_美化.docx'

# 颜色主题
PRIMARY_COLOR = RGBColor(46, 117, 182)      # 深蓝 #2E75B6
SECONDARY_COLOR = RGBColor(31, 78, 121)     # 深灰蓝 #1F4E79
ACCENT_COLOR = RGBColor(27, 161, 113)       # 青绿 #1BA171
TEXT_DARK = RGBColor(51, 51, 51)            # 深灰 #333333
TEXT_LIGHT = RGBColor(102, 102, 102)        # 浅灰 #666666
HEADER_BG = RGBColor(221, 232, 240)         # 淡蓝 #D5E8F0


def set_font(run, font_name='微软雅黑', size=Pt(11), bold=False, italic=False, color=None):
    """设置字体样式"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def add_heading(doc, text, level=1, color=None):
    """添加标题"""
    if color is None:
        color = PRIMARY_COLOR if level == 1 else SECONDARY_COLOR
    
    sizes = {1: Pt(18), 2: Pt(15), 3: Pt(13), 4: Pt(12)}
    size = sizes.get(level, Pt(12))
    
    heading = doc.add_heading(text, level=level)
    run = heading.runs[0]
    set_font(run, size=size, bold=True, color=color)
    
    # 添加下划线效果（对于一级标题）
    if level == 1:
        run.underline = True
    
    # 添加间距
    paragraph = heading.paragraph
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(18)
    paragraph_format.space_after = Pt(12)
    
    return heading


def add_subheading(doc, text, color=SECONDARY_COLOR):
    """添加副标题"""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    set_font(run, size=Pt(13), bold=True, color=color)
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(6)
    return paragraph


def add_text_paragraph(doc, text, style='normal', color=None, size=None, bold=False, left_indent=Pt(0)):
    """添加文本段落"""
    if color is None:
        color = TEXT_DARK
    if size is None:
        size = Pt(11)
    
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.left_indent = left_indent
    
    run = paragraph.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    return paragraph


def add_bulleted_list(doc, items, color=None):
    """添加项目符号列表"""
    if color is None:
        color = TEXT_DARK
    
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        set_font(run, size=Pt(10.5), color=color)


def add_numbered_list(doc, items, color=None):
    """添加编号列表"""
    if color is None:
        color = TEXT_DARK
    
    for item in items:
        p = doc.add_paragraph(style='List Number')
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        set_font(run, size=Pt(10.5), color=color)


def add_table(doc, headers, data, header_color=None):
    """添加表格"""
    if header_color is None:
        header_color = HEADER_BG
    
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.style = 'Table Grid'
    
    # 设置表头
    header_row = table.rows[0]
    header_cell = None
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        
        # 设置背景色
        tcPr = cell._tc.get_or_add_tcPr()
        shd = tcPr.find(qn('w:shd'))
        if shd is None:
            shd = qn('w:shd')
            tcPr.append(shd)
        shd.set(qn('w:fill'), header_color.rgb.hex if hasattr(header_color.rgb, 'hex') else 'D5E8F0')
        
        # 设置表头字体
        run = cell.paragraphs[0].runs[0]
        set_font(run, size=Pt(11), bold=True, color=SECONDARY_COLOR)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 设置数据行
    for i, row_data in enumerate(data):
        row = table.rows[i + 1]
        for j, value in enumerate(row_data):
            cell = row.cells[j]
            cell.text = value
            run = cell.paragraphs[0].runs[0]
            set_font(run, size=Pt(10.5), color=TEXT_DARK)
    
    # 设置表格宽度
    table.width = Inches(6.5)
    
    return table


def add_separator(doc):
    """添加分隔线"""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run('─' * 50)
    set_font(run, size=Pt(9), color=RGBColor(200, 200, 200))
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(12)


def parse_markdown(md_file):
    """解析 Markdown 文件"""
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    return content


def generate_resume(doc, md_content):
    """生成简历文档"""
    
    # 1. 基本信息部分
    lines = md_content.strip().split('\n')
    
    current_section = None
    bullet_items = []
    in_bullets = False
    
    # 姓名和联系方式（从第一行提取）
    if lines:
        # 姓名
        name_line = lines[0].strip()
        if name_line.startswith('#'):
            name_line = name_line.replace('#', '').strip()
        
        # 添加姓名作为大标题
        title_paragraph = doc.add_paragraph()
        title_run = title_paragraph.add_run(name_line)
        set_font(title_run, size=Pt(24), bold=True, color=PRIMARY_COLOR)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_paragraph.paragraph_format.space_after = Pt(12)
        
        # 职位（从第二行或从内容中提取）
        if len(lines) > 1 and 'AI 技术总监' in lines[1]:
            job_title = lines[1].strip().replace('-', '').replace('—', '').strip()
            if job_title:
                job_para = doc.add_paragraph()
                job_run = job_para.add_run(job_title)
                set_font(job_run, size=Pt(14), color=TEXT_LIGHT, italic=True)
                job_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                job_para.paragraph_format.space_after = Pt(18)
    
    # 解析其余内容
    i = 2  # 跳过姓名和职位行
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue
        
        # 一级标题
        if line.startswith('## '):
            section_title = line.replace('##', '').strip()
            add_heading(doc, section_title, level=1)
            i += 1
            continue
        
        # 二级标题
        elif line.startswith('### '):
            subsection_title = line.replace('###', '').strip()
            add_subheading(doc, subsection_title)
            i += 1
            continue
        
        # 项目符号列表项
        elif line.startswith('- ') or line.startswith('• '):
            bullet_text = line[2:].strip()
            bullet_items.append(bullet_text)
            i += 1
            continue
        
        # 普通文本
        else:
            # 如果有未处理的 bullet 项，先添加它们
            if bullet_items:
                add_bulleted_list(doc, bullet_items)
                bullet_items = []
            
            # 处理普通文本（可能是联系方式等）
            if ':' in line and not line.startswith('|'):
                # 键值对格式
                key, value = line.split(':', 1)
                key = key.strip()
                value = value.strip()
                
                para = doc.add_paragraph()
                para.paragraph_format.space_after = Pt(4)
                
                key_run = para.add_run(key + ':')
                set_font(key_run, size=Pt(11), bold=True, color=SECONDARY_COLOR)
                
                value_run = para.add_run(' ' + value)
                set_font(value_run, size=Pt(11), color=TEXT_DARK)
            else:
                add_text_paragraph(doc, line)
            
            i += 1
            continue
    
    # 处理剩余的 bullet 项
    if bullet_items:
        add_bulleted_list(doc, bullet_items)


def main():
    """主函数"""
    print("📄 正在读取 Markdown 文件...")
    md_content = parse_markdown(INPUT_MD)
    
    print("🎨 正在生成美化的 Word 文档...")
    doc = Document()
    
    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
    
    # 生成简历内容
    generate_resume(doc, md_content)
    
    # 保存文档
    doc.save(OUTPUT_DOCX)
    print(f"✅ 简历已保存至：{OUTPUT_DOCX}")
    print(f"📊 文件大小：{os.path.getsize(OUTPUT_DOCX) / 1024:.1f} KB")


if __name__ == '__main__':
    main()
