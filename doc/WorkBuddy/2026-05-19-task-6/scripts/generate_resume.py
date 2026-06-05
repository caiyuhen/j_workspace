#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
蔡宇衡 AI技术总监 优化版简历 - Word生成脚本
基于JD匹配分析，针对性优化内容并美化排版
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

OUTPUT_PATH = os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    "..", "deliverables", "hr-operations",
    "蔡宇衡-AI技术总监-优化简历-2026-05-19.docx"
)

# ── 颜色常量 ──
NAVY = RGBColor(0x1F, 0x38, 0x64)       # 深藏青 - 主标题/一级标题
DARK_BLUE = RGBColor(0x2E, 0x74, 0xB5)   # 中蓝 - 二级标题
ACCENT_BLUE = RGBColor(0x44, 0x72, 0xC4)  # 强调蓝 - 装饰线
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)    # 正文深灰
MED_GRAY = RGBColor(0x66, 0x66, 0x66)     # 次要信息
LIGHT_BG = RGBColor(0xF2, 0xF6, 0xFA)     # 浅蓝背景


def set_cell_shading(cell, color_hex):
    """设置单元格背景色"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_horizontal_line(doc, color="1F3864", thickness=6):
    """添加水平分割线"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="{thickness}" w:space="1" w:color="{color}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    return p


def set_run_font(run, cn_font="Microsoft YaHei", en_font="Calibri", size=10.5,
                 bold=False, color=None, italic=False):
    """设置run的字体属性"""
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    run.font.name = en_font
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), cn_font)
    rFonts.set(qn('w:ascii'), en_font)
    rFonts.set(qn('w:hAnsi'), en_font)


def add_section_title(doc, title, level=1):
    """添加带装饰的章节标题"""
    p = doc.add_paragraph()
    if level == 1:
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        # 添加左侧蓝色竖线装饰
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:left w:val="single" w:sz="24" w:space="8" w:color="1F3864"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)
        run = p.add_run(title)
        set_run_font(run, cn_font="Microsoft YaHei", size=14, bold=True, color=NAVY)
    elif level == 2:
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(title)
        set_run_font(run, cn_font="Microsoft YaHei", size=11, bold=True, color=DARK_BLUE)
    elif level == 3:
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(title)
        set_run_font(run, cn_font="Microsoft YaHei", size=10.5, bold=True, color=DARK_GRAY)
    return p


def add_body_text(doc, text, indent=False, bold_prefix=None, bullet=False):
    """添加正文段落"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = Pt(16)
    if indent:
        p.paragraph_format.first_line_indent = Pt(21)
    if bullet:
        p.paragraph_format.left_indent = Pt(14)
        prefix = "● "
    else:
        prefix = ""

    if bold_prefix:
        run_b = p.add_run(prefix + bold_prefix)
        set_run_font(run_b, size=10, bold=True, color=DARK_GRAY)
        run_t = p.add_run(text)
        set_run_font(run_t, size=10, color=DARK_GRAY)
    else:
        run = p.add_run(prefix + text)
        set_run_font(run, size=10, color=DARK_GRAY)
    return p


def add_bullet_item(doc, text, bold_prefix=None, level=0):
    """添加列表项"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = Pt(15)
    indent_level = 14 + level * 14
    p.paragraph_format.left_indent = Pt(indent_level)

    bullet_char = "▸" if level == 0 else "◦"
    if bold_prefix:
        run_b = p.add_run(f"{bullet_char} {bold_prefix}")
        set_run_font(run_b, size=9.5, bold=True, color=DARK_GRAY)
        run_t = p.add_run(text)
        set_run_font(run_t, size=9.5, color=DARK_GRAY)
    else:
        run = p.add_run(f"{bullet_char} {text}")
        set_run_font(run, size=9.5, color=DARK_GRAY)
    return p


def add_work_entry(doc, company, title, period, highlights):
    """添加工作经历条目"""
    # 公司名称 + 职位 + 时间
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = Pt(16)

    run_company = p.add_run(company)
    set_run_font(run_company, cn_font="Microsoft YaHei", size=10.5, bold=True, color=NAVY)

    run_sep = p.add_run("  |  ")
    set_run_font(run_sep, size=10.5, color=MED_GRAY)

    run_title = p.add_run(title)
    set_run_font(run_title, cn_font="Microsoft YaHei", size=10.5, bold=True, color=DARK_BLUE)

    # 时间右对齐 - 通过tab
    run_time = p.add_run(f"    {period}")
    set_run_font(run_time, size=9.5, color=MED_GRAY)

    # 工作亮点
    for item in highlights:
        if isinstance(item, tuple):
            add_bullet_item(doc, item[1], bold_prefix=item[0])
        else:
            add_bullet_item(doc, item)
    return p


def add_project_entry(doc, name, period, desc, tech_highlights, results):
    """添加项目经历条目"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = Pt(16)

    run_name = p.add_run(name)
    set_run_font(run_name, cn_font="Microsoft YaHei", size=10.5, bold=True, color=NAVY)

    run_time = p.add_run(f"    {period}")
    set_run_font(run_time, size=9.5, color=MED_GRAY)

    # 项目描述
    if desc:
        add_body_text(doc, desc)

    # 技术亮点
    for item in tech_highlights:
        if isinstance(item, tuple):
            add_bullet_item(doc, item[1], bold_prefix=item[0])
        else:
            add_bullet_item(doc, item)

    # 成果
    if results:
        for r in results:
            add_bullet_item(doc, r, bold_prefix="成果：")
    return p


def build_resume():
    doc = Document()

    # ── 页面设置 ──
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    # ── 默认字体 ──
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    style.font.color.rgb = DARK_GRAY
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # ════════════════════════════════════════════
    #  姓名 + 联系方式（顶部横排）
    # ════════════════════════════════════════════
    p_name = doc.add_paragraph()
    p_name.paragraph_format.space_before = Pt(0)
    p_name.paragraph_format.space_after = Pt(2)
    p_name.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_name = p_name.add_run("蔡宇衡")
    set_run_font(run_name, cn_font="Microsoft YaHei", en_font="Calibri", size=22, bold=True, color=NAVY)

    # 副标题
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(4)
    run_sub = p_sub.add_run("AI 技术总监 ｜ 22年研发经验 ｜ 15年技术管理")
    set_run_font(run_sub, cn_font="Microsoft YaHei", size=12, color=DARK_BLUE)

    # 联系方式 - 用表格横排
    tbl = doc.add_table(rows=1, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.autofit = True

    contact_info = [
        ("📞", "13810357924"),
        ("📧", "caiyuheng81@outlook.com"),
        ("📍", "北京"),
        ("🎓", "硕士 · 山东大学"),
    ]
    for i, (icon, text) in enumerate(contact_info):
        cell = tbl.cell(0, i)
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run_i = p.add_run(icon + " ")
        set_run_font(run_i, size=9, color=MED_GRAY)
        run_t = p.add_run(text)
        set_run_font(run_t, size=9, color=DARK_GRAY)
        # 移除单元格边框
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'</w:tcBorders>'
        )
        tcPr.append(tcBorders)

    # 顶部装饰线
    add_horizontal_line(doc, color="1F3864", thickness=8)

    # ════════════════════════════════════════════
    #  求职意向（精简）
    # ════════════════════════════════════════════
    add_section_title(doc, "求职意向")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    items = [
        ("目标职位：", "AI 技术总监 / 技术 VP"),
        ("期望行业：", "人工智能 · 医疗健康 · SaaS 平台"),
        ("工作地点：", "北京"),
    ]
    for i, (label, val) in enumerate(items):
        if i > 0:
            run_sep = p.add_run("    ")
            set_run_font(run_sep, size=9.5, color=MED_GRAY)
        run_l = p.add_run(label)
        set_run_font(run_l, size=9.5, bold=True, color=NAVY)
        run_v = p.add_run(val)
        set_run_font(run_v, size=9.5, color=DARK_GRAY)

    add_horizontal_line(doc, color="4472C4", thickness=4)

    # ════════════════════════════════════════════
    #  核心优势（4大模块 - 用表格+背景色卡片式）
    # ════════════════════════════════════════════
    add_section_title(doc, "核心优势")

    advantages = [
        ("🏗 技术战略与架构规划", [
            "22年软件开发经验，15年技术管理背景，擅长制定中长期技术演进路线图",
            "主导系统架构升级：前端框架迁移（jQuery→Vue3微前端）、双API网关解耦、BFF层建设，发布频率从月级提升至周级",
            "精通微服务架构设计，能在关键技术问题上做兜底决策",
            "具备从0到1组建百人级产研团队经验（心医国际/圣方时期），对系统连续性、AI落地、研发交付节奏负总责",
        ]),
        ("🤖 大模型应用与技术落地", [
            "深度实践LLM应用开发，精通Prompt Engineering、RAG检索增强、Agent编排、Function Calling、Skill框架设计",
            "独立交付AI Agent系统上线（药物警戒智能体），完成从架构设计到生产部署的全流程",
            "熟悉LangChain、LlamaIndex、Dify、AutoGen等主流LLM应用框架",
            "主导构建企业级RAG知识库系统，实现文档分词、向量存储、语义检索的全链路技术闭环",
        ]),
        ("🏥 医疗行业与SaaS经验", [
            "15年深耕医疗健康领域，熟悉ICH GCP、FDA 21 CFR Part 11、CDISC/SDTM等合规标准",
            "主导开发CTMS+EDC临床试验管理系统，具备私有部署+多租户SaaS系统完整架构经验",
            "熟悉企业微信应用生态，有医疗行业政企SaaS产品落地经验",
            "主导10+款AI医疗产品从原型设计到NMPA注册交付全周期管理",
        ]),
        ("🤝 跨组织协作与团队管理", [
            "在圣方医药曾同时协调药监局合规团队、临床科室、IT运维三方推进系统落地，保障各方需求的同时守住团队迭代节奏和架构决策主导权",
            "管理外包运维伙伴服务质量与SLA，确保系统可用性达99.9%",
            "牵头40+真实世界研究项目，主导国家级课题申报及行业标准制定（NMPA/CDE/北京大学合作）",
        ]),
    ]

    for title, items in advantages:
        add_section_title(doc, title, level=2)
        for item in items:
            add_bullet_item(doc, item)

    add_horizontal_line(doc, color="4472C4", thickness=4)

    # ════════════════════════════════════════════
    #  工作经历（优化版）
    # ════════════════════════════════════════════
    add_section_title(doc, "工作经历")

    # 1. 安顿健康 - 修改职位名称说明
    add_work_entry(doc,
        "安顿健康科技有限公司",
        "AI技术负责人（高级算法工程师 · 初创公司兼任技术规划与团队管理）",
        "2025.07 - 至今",
        [
            ("技术战略：", "制定AI技术中长期发展规划，主导跨部门技术协作机制建设，规划技术演进路线图"),
            ("系统架构：", "负责智能穿戴设备健康预测系统架构设计，管理外包开发团队SLA交付质量与系统连续性"),
            ("中医智诊机器人：", "基于LangChain+RAG+AutoGen构建多Agent协作框架，设计Skill框架扩展机制；构建8万+方剂、2000+证候、9000+中药材的RAG知识库；结合望诊图像识别、切诊传感器数据、问诊智能问答实现四诊法数字化——系统已上线运行"),
            ("大数据平台：", "基于Hadoop/Spark构建高并发数据处理平台，设计数据标准化SOP与安全SOP"),
        ]
    )

    # 2. 圣方医药
    add_work_entry(doc,
        "圣方(上海)医药研发有限公司",
        "高级算法工程师 / 技术负责人",
        "2021.05 - 2025.04",
        [
            ("技术路线：", "制定AI技术演进路线图，主导跨部门研发管理体系建设，推动从算法研发到产品注册的完整闭环"),
            ("大模型方向：", "主导大模型应用技术决策（选型/架构/Skill框架/Agent编排/RAG知识库），优化医学影像大模型架构，推动多模态融合研究"),
            ("产品落地：", "主导10+款AI医疗产品从原型设计到NMPA注册交付的全周期管理"),
            ("团队建设：", "搭建高绩效研发团队（约30人，含算法/后端/前端/测试），建立技术人才选拔培养机制"),
            ("行业标准：", "牵头与国内外机构合作研究，参与制定3项行业标准"),
        ]
    )

    # 3. 心医国际 - 大幅扩写
    add_work_entry(doc,
        "心医国际",
        "副总裁 / 技术VP",
        "2020.09 - 2021.05",
        [
            ("团队管理：", "统筹产研团队约60人，协调前端/后端/AI/数据/运维多条线，建立跨团队sprint机制"),
            ("架构升级：", "主导SaaS平台从单体架构向私有部署+多租户架构迁移——引入BFF层解耦双API网关，前端从传统框架升级至组件化架构，系统可用性提升至99.9%"),
            ("AI落地：", "牵头CDSS知识图谱系统落地，构建基于知识图谱的临床决策支持系统，服务100+医院；开发联邦学习平台，实现隐私保护下的多中心协作"),
            ("区域协同：", "陕西/贵州省级项目支持100+医院急救体系协同数据交互；药物警戒系统在300+医院分布式部署"),
        ]
    )

    # 4. 首佑医学
    add_work_entry(doc,
        "北京首佑医学科技",
        "大数据人工智能总监",
        "2018.09 - 2020.09",
        [
            ("知识图谱：", "通过NLP技术解析150万+电子病历构建精神疾病知识图谱，支撑京津冀三地医疗数据平台建设"),
            ("核心算法：", "主导语音识别（CNN/RNN）与MRI影像分析模型研发，支撑国自然课题"),
            ("产品落地：", "设计双抗药物浓度动态监测算法，完成6医院临床落地"),
            ("行业标准：", "AI辅助诊断框架技术方案入选3个省级医疗标准"),
        ]
    )

    # 5. 安华亿能
    add_work_entry(doc,
        "安华亿能医疗影像科技",
        "CTO",
        "2012.04 - 2018.09",
        [
            ("技术创新：", "首创全球颈动脉三维超声诊断系统，斑块自动检测准确率98.2%"),
            ("SaaS平台：", "开发云诊所影像SaaS平台，完成130+家三甲医院PACS系统无缝对接"),
            ("系统重构：", "主导平台从单体架构向微服务架构改造，引入分布式存储与计算框架"),
            ("医疗器械：", "主导II类医疗器械从概念到NMPA认证全流程，建立设计控制闭环"),
        ]
    )

    # 6. 乾镛科技 - 注明并行
    add_work_entry(doc,
        "乾镛科技（上海）| CEO（与安华亿能同期创业项目）",
        "CEO",
        "2012.04 - 2018.09",
        [
            "打造四大智能平台（中国卒中学院APP/话说卒中APP/心电云/影像会诊平台），年直播学术会议120+场，数字化管理病历250万份",
            "获\"十三五\"国家重点研发计划支持（课题编号：2018YFC1314700）",
        ]
    )

    # 7-9. 早期经历 - 压缩
    add_section_title(doc, "早期经历", level=2)
    early_items = [
        "长城国际体育传播 | 总裁（2010-2012）：主导国际顶级赛事中国区运营，开创\"体育+科技\"商业模式",
        "北京新传德国际版权交易中心 | 副总裁（2008-2010）：组建40人技术研发团队，主导版权交易系统架构升级",
        "清华同方 | 架构师（2006-2008）：主导ERP/CRM系统全生命周期开发，制定基于微服务/SOA的技术选型方案",
        "三星中国 | 架构师（2004-2006）：主导企业级系统架构设计，设计电算财务系统（年处理资金60亿美元）",
        "SOHU | 开发工程师（2002-2004）：NLP算法在商业场景的工程化落地",
    ]
    for item in early_items:
        add_body_text(doc, item)

    add_horizontal_line(doc, color="4472C4", thickness=4)

    # ════════════════════════════════════════════
    #  核心项目经验（精简，删除智能驾驶）
    # ════════════════════════════════════════════
    add_section_title(doc, "核心项目经验")

    add_project_entry(doc,
        "AI Agent系统 — 药物警戒智能体",
        "2024.01 - 2024.11",
        "构建基于大模型的药物警戒系统，实现AE/SAE自动采集、分级、多语言翻译、全球分发",
        [
            ("Agent框架：", "LangChain（工具集成）+ AutoGen（多Agent协作）+ Dify（工作流编排）"),
            ("RAG检索：", "基于LlamaIndex实现定向检索，Zep向量存储+MemGPT记忆逻辑控制"),
            ("决策框架：", "多LLM支持（GPT/Claude/Gemini/DeepSeek/Doubao）+ PDDL规划算法"),
            ("提示词工程：", "PromptPerfect自动优化+BERT语义理解偏差修正"),
        ],
        ["系统已上线运行，实现从数据采集到报告生成的全流程自动化"]
    )

    add_project_entry(doc,
        "CTMS+EDC 临床试验管理系统",
        "进行中",
        "主导设计临床试验管理系统，符合ICH GCP、FDA 21 CFR Part 11、CDISC/SDTM标准",
        [
            ("后端：", "Node.js/TypeScript + Express.js + Prisma ORM + PostgreSQL，后端30+模块"),
            ("前端：", "React + Vite + Ant Design + Zustand，前端21页面已完成"),
            ("核心模块：", "EdcTemplate、CrfForm、CrfData、AdverseEvent、角色权限、全流程审批"),
            ("数据合规：", "6类CSV/JSON导出（CDISC标准eCRF与SDTM导出），私有部署+多租户架构"),
        ],
        ["正在推进RAG检索系统构建（文档分词+向量库导入），调研CDISC/SDTM标准合规"]
    )

    add_project_entry(doc,
        "健康管理与医疗保险结合系统",
        "2024.01 - 2024.11",
        "构建覆盖核保、核赔、医疗网络管理、支付结算全流程的商保平台",
        [
            ("核赔审核：", "AI算法实现医疗费用合理性评估、过度医疗识别"),
            ("数据治理：", "ICD-10编码映射、非结构化病历解析、DICOM标准接口"),
            ("智能监测：", "慢病管理、欺诈风险监测、疾病发生率预测图谱"),
        ],
        ["标准化接口与200+医疗机构系统直连，日均处理30万+医疗单据，审核准确率98.6%"]
    )

    add_horizontal_line(doc, color="4472C4", thickness=4)

    # ════════════════════════════════════════════
    #  技术能力矩阵（展开！不再写"略"）
    # ════════════════════════════════════════════
    add_section_title(doc, "技术能力矩阵")

    # 用表格展示，更清晰
    tech_data = [
        ("大模型应用", "Prompt Engineering · RAG · Agent编排 · Function Calling · Skill框架\nLangChain · LlamaIndex · Dify · AutoGen · Zep · MemGPT"),
        ("后端架构\n（主力）", "Node.js/TypeScript（BFF层开发/微服务）\nJava（Spring Boot，历史系统维护）\nPython（FastAPI/Flask，AI服务与数据处理）\nGo（高并发网关服务）"),
        ("前端技术", "React · Vue3/微前端 · Vite · Ant Design · Zustand"),
        ("数据架构", "PostgreSQL · MongoDB · Milvus（向量库） · Redis · Kafka · Spark · Hadoop"),
        ("AI/算法", "Transformer · PyTorch · TensorFlow · 知识图谱 · 联邦学习 · 多模态融合"),
        ("云原生", "Docker · Kubernetes · Jenkins · HPA · 边缘计算K3S"),
        ("医疗合规", "ICH GCP E6(R2) · FDA 21 CFR Part 11 · CDISC/SDTM · GDPR · OMOP"),
        ("企业应用", "企业微信API · 微信生态 · 政企SaaS · 私有部署 · 多租户系统"),
    ]

    tbl_tech = doc.add_table(rows=len(tech_data) + 1, cols=2)
    tbl_tech.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_tech.autofit = True

    # 表头
    hdr = tbl_tech.rows[0]
    for i, text in enumerate(["能力维度", "核心技术栈"]):
        cell = hdr.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_run_font(run, cn_font="Microsoft YaHei", size=9.5, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        set_cell_shading(cell, "1F3864")

    # 数据行
    for idx, (dim, techs) in enumerate(tech_data):
        row = tbl_tech.rows[idx + 1]
        # 维度
        cell_d = row.cells[0]
        cell_d.text = ""
        p = cell_d.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_d = p.add_run(dim)
        set_run_font(run_d, cn_font="Microsoft YaHei", size=9, bold=True, color=NAVY)
        set_cell_shading(cell_d, "F2F6FA")

        # 技术栈
        cell_t = row.cells[1]
        cell_t.text = ""
        p = cell_t.paragraphs[0]
        run_t = p.add_run(techs)
        set_run_font(run_t, size=9, color=DARK_GRAY)
        if idx % 2 == 1:
            set_cell_shading(cell_t, "F8F8F8")

    # 表格样式 - 去掉默认边框，只留三线表效果
    tbl_tech.style = doc.styles['Table Grid']

    add_horizontal_line(doc, color="4472C4", thickness=4)

    # ════════════════════════════════════════════
    #  教育背景
    # ════════════════════════════════════════════
    add_section_title(doc, "教育背景")

    edu_items = [
        ("山东大学", "硕士 · 数据科学与大数据技术", "2004 - 2007"),
        ("北京科技大学", "本科 · 计算机应用", "1998 - 2002"),
    ]
    for school, degree, period in edu_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run_s = p.add_run(school)
        set_run_font(run_s, cn_font="Microsoft YaHei", size=10, bold=True, color=NAVY)
        run_d = p.add_run(f"  {degree}")
        set_run_font(run_d, size=10, color=DARK_GRAY)
        run_p = p.add_run(f"    {period}")
        set_run_font(run_p, size=9.5, color=MED_GRAY)

    add_horizontal_line(doc, color="4472C4", thickness=4)

    # ════════════════════════════════════════════
    #  资格证书 & 附加说明
    # ════════════════════════════════════════════
    add_section_title(doc, "资格证书与附加说明")

    cert_items = [
        "高级人工智能训练师",
        "高级健康管理师",
        "大学英语四级（读写精通），专业四级",
    ]
    for item in cert_items:
        add_bullet_item(doc, item)

    add_section_title(doc, "关键加分项", level=2)
    bonus = [
        "15年医疗/健康/CRO领域深耕经验，主导政企SaaS产品落地",
        "熟悉LangChain/LlamaIndex/Dify等主流LLM应用框架，独立交付AI Agent系统上线",
        "具备私有部署+多租户SaaS系统完整架构经验（心医国际/安华亿能）",
        "主导过多次老系统重构（心医国际前端框架升级/微服务改造/安华亿能数据中台建设）",
        "熟悉企业微信应用生态，有CTMS+EDC系统集成经验",
    ]
    for item in bonus:
        add_bullet_item(doc, item)

    add_section_title(doc, "知识产权", level=2)
    ip_items = [
        "3+项发明专利（含智能驾驶感知融合、AI辅助诊断等方向）",
        "参与制定3+项行业标准，技术方案入选3+个省级标准",
        "主导国自然课题、\"十三五\"国家重点研发计划",
    ]
    for item in ip_items:
        add_bullet_item(doc, item)

    # ── 页脚（页码） ──
    footer = section.footer
    footer.is_linked_to_previous = False
    p_footer = footer.paragraphs[0]
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加 "第 X 页" 页码
    run1 = p_footer.add_run("— ")
    set_run_font(run1, size=8, color=MED_GRAY)
    # PAGE field
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run_field = p_footer.add_run()
    run_field._r.append(fldChar1)
    run_instr = p_footer.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run_instr._r.append(instrText)
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run_sep2 = p_footer.add_run()
    run_sep2._r.append(fldChar2)
    run_num = p_footer.add_run("1")
    set_run_font(run_num, size=8, color=MED_GRAY)
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run_end = p_footer.add_run()
    run_end._r.append(fldChar3)
    run2 = p_footer.add_run(" —")
    set_run_font(run2, size=8, color=MED_GRAY)

    # ── 保存 ──
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"简历已生成：{OUTPUT_PATH}")
    return OUTPUT_PATH


if __name__ == "__main__":
    path = build_resume()
    print(f"完成！文件：{path}")
