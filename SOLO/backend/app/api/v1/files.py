"""
附件上传与阅读理解 API。

说明：LLM 服务本身已经内置 RAG，本模块不重复建设 RAG 或向量库，
只负责上传文件、提取附件文本，并把附件内容作为上下文交给现有 LLM 服务理解。
"""
from __future__ import annotations

import os
import uuid
from datetime import datetime
from pathlib import Path
from typing import Dict, Optional

from fastapi import APIRouter, Depends, File, HTTPException, UploadFile
from fastapi.responses import FileResponse
from pydantic import BaseModel, Field

from app.api.v1.auth import TokenData, get_current_active_user
from app.services.llm_service import llm_service


router = APIRouter()

UPLOAD_DIR = Path(os.getenv("UPLOAD_DIR", "uploads")).resolve()
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

MAX_UPLOAD_SIZE = 20 * 1024 * 1024
MAX_CONTEXT_CHARS = 24000
ALLOWED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}

_files_db: Dict[str, Dict] = {}


class FileUploadResponse(BaseModel):
    file_id: str
    filename: str
    content_type: Optional[str] = None
    char_count: int
    text_preview: str
    created_at: datetime


class FileInfoResponse(FileUploadResponse):
    owner_user_id: str


class FileAskRequest(BaseModel):
    question: str = Field(..., min_length=1, description="针对附件提出的问题")


class FileAskResponse(BaseModel):
    file_id: str
    filename: str
    question: str
    answer: str


def _safe_filename(filename: str) -> str:
    cleaned = Path(filename or "uploaded_file").name
    return cleaned.replace("\x00", "") or "uploaded_file"


def _decode_text(raw: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "gbk"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="ignore")


def _extract_txt_or_md(path: Path) -> str:
    return _decode_text(path.read_bytes())


def _extract_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except Exception as exc:
        raise HTTPException(status_code=500, detail="服务端缺少 PDF 解析依赖 pypdf") from exc

    reader = PdfReader(str(path))
    parts = []
    for page in reader.pages:
        parts.append(page.extract_text() or "")
    text = "\n".join(parts).strip()
    if not text:
        raise HTTPException(status_code=400, detail="未能从 PDF 中提取文本，可能是扫描版图片 PDF，需要 OCR")
    return text


def _extract_docx(path: Path) -> str:
    try:
        from docx import Document
    except Exception as exc:
        raise HTTPException(status_code=500, detail="服务端缺少 DOCX 解析依赖 python-docx") from exc

    document = Document(str(path))
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text and p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    text = "\n".join(paragraphs).strip()
    if not text:
        raise HTTPException(status_code=400, detail="未能从 DOCX 中提取文本")
    return text


def extract_text(path: Path, extension: str) -> str:
    if extension in {".txt", ".md"}:
        return _extract_txt_or_md(path)
    if extension == ".pdf":
        return _extract_pdf(path)
    if extension == ".docx":
        return _extract_docx(path)
    raise HTTPException(status_code=400, detail=f"不支持的文件类型: {extension}")


def _get_owned_file(file_id: str, user_id: str) -> Dict:
    record = _files_db.get(file_id)
    if not record:
        raise HTTPException(status_code=404, detail="附件不存在")
    if record.get("owner_user_id") != user_id:
        raise HTTPException(status_code=403, detail="无权访问此附件")
    return record


@router.post("/upload", response_model=FileUploadResponse)
async def upload_file(
    file: UploadFile = File(...),
    current_user: TokenData = Depends(get_current_active_user),
):
    filename = _safe_filename(file.filename or "uploaded_file")
    extension = Path(filename).suffix.lower()
    if extension not in ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=400, detail=f"不支持的文件类型: {extension or '无扩展名'}")

    raw = await file.read()
    if not raw:
        raise HTTPException(status_code=400, detail="上传文件为空")
    if len(raw) > MAX_UPLOAD_SIZE:
        raise HTTPException(status_code=400, detail="上传文件超过 20MB 限制")

    file_id = str(uuid.uuid4())
    stored_name = f"{file_id}{extension}"
    stored_path = UPLOAD_DIR / stored_name
    stored_path.write_bytes(raw)

    text = extract_text(stored_path, extension).strip()
    if not text:
        raise HTTPException(status_code=400, detail="未能从附件中提取文本")

    now = datetime.now()
    record = {
        "file_id": file_id,
        "filename": filename,
        "content_type": file.content_type,
        "path": str(stored_path),
        "text": text,
        "char_count": len(text),
        "created_at": now,
        "owner_user_id": current_user.user_id,
    }
    _files_db[file_id] = record

    return FileUploadResponse(
        file_id=file_id,
        filename=filename,
        content_type=file.content_type,
        char_count=len(text),
        text_preview=text[:500],
        created_at=now,
    )


@router.get("/{file_id}/download")
async def download_file(
    file_id: str,
    current_user: TokenData = Depends(get_current_active_user),
):
    """下载已上传附件。"""
    record = _get_owned_file(file_id, current_user.user_id)
    stored_path = Path(record["path"])
    if not stored_path.exists() or not stored_path.is_file():
        raise HTTPException(status_code=404, detail="附件文件不存在或已被删除")

    return FileResponse(
        path=str(stored_path),
        media_type=record.get("content_type") or "application/octet-stream",
        filename=record["filename"],
    )


@router.get("/{file_id}", response_model=FileInfoResponse)
async def get_file_info(
    file_id: str,
    current_user: TokenData = Depends(get_current_active_user),
):
    record = _get_owned_file(file_id, current_user.user_id)
    return FileInfoResponse(
        file_id=record["file_id"],
        filename=record["filename"],
        content_type=record.get("content_type"),
        char_count=record["char_count"],
        text_preview=record["text"][:500],
        created_at=record["created_at"],
        owner_user_id=record["owner_user_id"],
    )


@router.post("/{file_id}/ask", response_model=FileAskResponse)
async def ask_file(
    file_id: str,
    request: FileAskRequest,
    current_user: TokenData = Depends(get_current_active_user),
):
    record = _get_owned_file(file_id, current_user.user_id)
    text = record["text"]
    context = text[:MAX_CONTEXT_CHARS]
    truncated_note = "\n\n注意：附件内容较长，以下仅包含前半部分文本。" if len(text) > MAX_CONTEXT_CHARS else ""

    prompt = f"""你是医学文档阅读理解助手。LLM 服务已经内置 RAG 能力，但本次任务应优先基于用户上传附件内容回答，不要编造附件中不存在的信息。

附件名称：{record['filename']}
附件内容：
{context}{truncated_note}

用户问题：{request.question}

请用中文回答。若附件信息不足，请明确说明不足之处。"""

    result = await llm_service.chat([
        {"role": "user", "content": prompt}
    ], session_id=f"file-{file_id}", max_tokens=4096)
    answer = result.get("content", "") if isinstance(result, dict) else str(result)

    return FileAskResponse(
        file_id=file_id,
        filename=record["filename"],
        question=request.question,
        answer=answer,
    )
